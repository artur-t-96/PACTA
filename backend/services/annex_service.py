"""
Annex service — generate contract annexes (aneksy)
Common annex types:
1. Rate change (zmiana stawki)
2. Role change (zmiana stanowiska)
3. Client change (zmiana klienta/projektu)
4. Address change (zmiana adresu firmy)
5. Termination agreement (porozumienie o rozwiązaniu)
6. Extension (przedłużenie umowy)
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.getenv("OUTPUT_DIR", "./output/contracts/annexes")


def generate_annex(
    contract_number: str,
    contractor_name: str,
    contractor_company: str,
    contractor_nip: str,
    contractor_address: str,
    annex_type: str,
    changes: list,  # [{"field": "stawka", "old": "150", "new": "180"}, ...]
    effective_date: str = "",
    notes: str = "",
) -> str:
    """Generate annex DOCX file."""
    
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
    
    today = datetime.now()
    today_str = today.strftime("%d.%m.%Y")
    effective = effective_date or today_str
    
    # Count existing annexes for this contract
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    existing = [f for f in os.listdir(OUTPUT_DIR) if contract_number.replace("/", "_") in f]
    annex_num = len(existing) + 1
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"ANEKS NR {annex_num}")
    run.bold = True
    run.font.size = Pt(14)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"do Umowy o współpracę B2B nr {contract_number}")
    run.font.size = Pt(11)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"z dnia {today_str}")
    run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Parties
    doc.add_paragraph("zawarty pomiędzy:")
    
    p = doc.add_paragraph()
    p.add_run("B2BNET S.A.").bold = True
    p.add_run(" z siedzibą w Warszawie, Al. Jerozolimskie 180, 02-486 Warszawa, "
              "wpisaną do rejestru przedsiębiorców KRS pod numerem 0000387063, "
              'NIP: 5711707392, reprezentowana przez Pana Artura Twardowskiego \u2014 Prezesa Zarzadu, '
              'zwana dalej \u201eZamawiajacym\u201d')
    
    doc.add_paragraph("a")
    
    p = doc.add_paragraph()
    p.add_run(f"{contractor_name}").bold = True
    nip_text = f", NIP: {contractor_nip}" if contractor_nip else ""
    p.add_run(f" prowadzacym/ca dzialalnosc gospodarcza pod firma: {contractor_company}"
              f", z adresem: {contractor_address}{nip_text}"
              ', zwanym dalej \u201ePartnerem\u201d')
    
    doc.add_paragraph()
    
    # Type-specific content
    type_titles = {
        "rate_change": "Zmiana wynagrodzenia",
        "role_change": "Zmiana zakresu usług",
        "client_change": "Zmiana klienta/projektu",
        "address_change": "Zmiana danych firmy",
        "termination": "Rozwiązanie umowy za porozumieniem stron",
        "extension": "Przedłużenie umowy",
        "other": "Zmiana warunków umowy",
    }
    
    # §1
    p = doc.add_paragraph()
    p.add_run("§ 1").bold = True
    
    doc.add_paragraph(
        f"Strony postanawiają wprowadzić następujące zmiany do Umowy o współpracę B2B "
        f"nr {contract_number}, ze skutkiem od dnia {effective}:"
    )
    
    # Changes
    for i, change in enumerate(changes, 1):
        field = change.get("field", "")
        old_val = change.get("old", "")
        new_val = change.get("new", "")
        description = change.get("description", "")
        
        if description:
            doc.add_paragraph(f"{i}. {description}", style="List Number")
        elif old_val and new_val:
            field_labels = {
                "stawka": "stawkę godzinową",
                "rola": "zakres świadczonych usług (stanowisko)",
                "klient": "klienta projektu",
                "adres": "adres prowadzenia działalności",
                "data_zakonczenia": "termin obowiązywania umowy",
            }
            label = field_labels.get(field, field)
            doc.add_paragraph(
                f'{i}. Strony zmieniaja {label} z: \u201e{old_val}\u201d na: \u201e{new_val}\u201d.',
                style="List Number"
            )
    
    doc.add_paragraph()
    
    # §2
    p = doc.add_paragraph()
    p.add_run("§ 2").bold = True
    doc.add_paragraph(
        "Pozostałe postanowienia Umowy nie ulegają zmianie i pozostają w mocy."
    )
    
    doc.add_paragraph()
    
    # §3
    p = doc.add_paragraph()
    p.add_run("§ 3").bold = True
    doc.add_paragraph(
        "Aneks sporządzono w dwóch jednobrzmiących egzemplarzach, po jednym dla każdej ze Stron."
    )
    
    if notes:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Uwagi: ").bold = True
        p.add_run(notes)
    
    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "________________________"
    table.cell(0, 1).text = "________________________"
    table.cell(1, 0).text = "Zamawiający (B2B.net S.A.)"
    table.cell(1, 1).text = f"Partner ({contractor_name})"
    
    # Save
    safe_number = contract_number.replace("/", "_")
    filename = f"Aneks_{annex_num}_{safe_number}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(output_path)
    
    return output_path
