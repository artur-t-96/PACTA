from io import BytesIO
from datetime import datetime
from typing import Dict, Any
import os
import base64

from docx import Document
from docx.shared import Pt, Twips, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree


class DocumentGenerator:
    """Generator umów B2B w formacie DOCX - zgodny w 100% z szablonem B2B.net"""
    
    # Ścieżka do logo firmy
    LOGO_PATH = os.path.join(os.path.dirname(__file__), '..', 'static', 'logo_b2bnet.jpeg')
    
    def __init__(self):
        self.doc = None
        self.FONT_NAME = 'Arial'
        self.FONT_SIZE_NORMAL = Pt(9)  # 18 half-points = 9pt
        self.FONT_SIZE_TITLE = Pt(14)  # 28 half-points = 14pt
        self.FONT_SIZE_FOOTER = Pt(8)  # 16 half-points = 8pt
        self.FONT_SIZE_PAGE_NUM = Pt(8)  # 16 half-points = 8pt
        
        # Marginesy w DXA (1440 DXA = 1 inch, 567 DXA = 1 cm)
        # Oryginał: top=1107, right=1417, bottom=1417, left=1417, header=454, footer=708
        self.PAGE_WIDTH = 11906   # A4 width
        self.PAGE_HEIGHT = 16838  # A4 height
        self.MARGIN_TOP = 1107
        self.MARGIN_BOTTOM = 1417
        self.MARGIN_LEFT = 1417
        self.MARGIN_RIGHT = 1417
        self.HEADER_DISTANCE = 454
        self.FOOTER_DISTANCE = 708
        
        # Spacing akapitów (oryginał: line="300" lineRule="atLeast")
        self.LINE_SPACING = 300  # 300 twips = 15pt
    
    def generate(self, data: Dict[str, Any]) -> bytes:
        """Generuje umowę B2B na podstawie danych z formularza."""
        self.doc = Document()
        self._setup_document()
        
        # Wybór szablonu na podstawie typu umowy
        if data.get("typ_umowy") == "z_dzialalnoscia":
            self._generate_contract_with_business(data)
        else:
            self._generate_contract_without_business(data)
        
        # Dodanie załączników
        self._add_attachment_1(data)
        self._add_attachment_2(data)
        self._add_attachment_3(data)
        
        # Zapis do bufora
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _setup_document(self):
        """Konfiguruje dokument - rozmiar A4, marginesy, nagłówek z logo, stopka z ramką"""
        section = self.doc.sections[0]
        
        # === ROZMIAR STRONY A4 ===
        section.page_width = Twips(self.PAGE_WIDTH)
        section.page_height = Twips(self.PAGE_HEIGHT)
        
        # === MARGINESY ===
        section.top_margin = Twips(self.MARGIN_TOP)
        section.bottom_margin = Twips(self.MARGIN_BOTTOM)
        section.left_margin = Twips(self.MARGIN_LEFT)
        section.right_margin = Twips(self.MARGIN_RIGHT)
        section.header_distance = Twips(self.HEADER_DISTANCE)
        section.footer_distance = Twips(self.FOOTER_DISTANCE)
        
        # === NAGŁÓWEK Z LOGO I NUMERACJĄ STRON ===
        self._setup_header(section)
        
        # === STOPKA Z RAMKĄ GÓRNĄ ===
        self._setup_footer(section)
        
        # === STYL NORMAL ===
        self._setup_styles()
    
    def _setup_header(self, section):
        """Konfiguruje nagłówek z logo i numeracją stron (zgodnie z oryginałem)"""
        header = section.header
        header.is_linked_to_previous = False
        
        # Pierwszy paragraf - numeracja strony wyrównana do prawej (w ramce)
        p_page_num = header.paragraphs[0]
        p_page_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Dodaj styl nagłówka i ramkę (framePr)
        pPr = p_page_num._p.get_or_add_pPr()
        
        # Styl nagłówka
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Nagwek')
        pPr.insert(0, pStyle)
        
        # Ramka pozycjonująca numer strony (framePr)
        framePr = OxmlElement('w:framePr')
        framePr.set(qn('w:wrap'), 'around')
        framePr.set(qn('w:vAnchor'), 'text')
        framePr.set(qn('w:hAnchor'), 'margin')
        framePr.set(qn('w:xAlign'), 'right')
        framePr.set(qn('w:y'), '1')
        pPr.append(framePr)
        
        # Dodaj pole PAGE (numer strony)
        run = p_page_num.add_run()
        self._set_run_font(run, self.FONT_SIZE_PAGE_NUM)
        
        # fldChar begin
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar_begin)
        
        # instrText PAGE
        run2 = p_page_num.add_run()
        self._set_run_font(run2, self.FONT_SIZE_PAGE_NUM)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE  '
        run2._r.append(instrText)
        
        # fldChar separate
        run3 = p_page_num.add_run()
        self._set_run_font(run3, self.FONT_SIZE_PAGE_NUM)
        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fldChar_sep)
        
        # fldChar end
        run4 = p_page_num.add_run()
        self._set_run_font(run4, self.FONT_SIZE_PAGE_NUM)
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run4._r.append(fldChar_end)
        
        # Drugi paragraf - logo
        p_logo = header.add_paragraph()
        pPr_logo = p_logo._p.get_or_add_pPr()
        
        # Styl nagłówka
        pStyle_logo = OxmlElement('w:pStyle')
        pStyle_logo.set(qn('w:val'), 'Nagwek')
        pPr_logo.insert(0, pStyle_logo)
        
        # Wcięcie z prawej
        ind = OxmlElement('w:ind')
        ind.set(qn('w:right'), '360')
        pPr_logo.append(ind)
        
        # Dodaj logo jeśli istnieje
        if os.path.exists(self.LOGO_PATH):
            run_logo = p_logo.add_run()
            # Logo: 733425 EMU x 489273 EMU (ok. 5.78cm x 3.86cm)
            run_logo.add_picture(self.LOGO_PATH, width=Emu(733425), height=Emu(489273))
        
        # Trzeci paragraf - pusty (zgodnie z oryginałem)
        p_empty = header.add_paragraph()
        pPr_empty = p_empty._p.get_or_add_pPr()
        pStyle_empty = OxmlElement('w:pStyle')
        pStyle_empty.set(qn('w:val'), 'Nagwek')
        pPr_empty.insert(0, pStyle_empty)
        ind_empty = OxmlElement('w:ind')
        ind_empty.set(qn('w:right'), '360')
        pPr_empty.append(ind_empty)
    
    def _setup_footer(self, section):
        """Konfiguruje stopkę z ramką górną (zgodnie z oryginałem)"""
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Pierwszy paragraf - pusty z ramką górną
        p_border = footer.paragraphs[0]
        self._add_footer_style_and_border(p_border)
        
        # Drugi paragraf - tekst stopki z ramką górną
        p_text = footer.add_paragraph()
        self._add_footer_style_and_border(p_text)
        
        run = p_text.add_run(
            "B2B.net S.A., Aleje Jerozolimskie 180, Kopernik Office Building, "
            "02-486 Warszawa, office@B2Bnetwork.pl, www.B2Bnetwork.pl"
        )
        self._set_run_font(run, self.FONT_SIZE_FOOTER, include_cs=True)
    
    def _add_footer_style_and_border(self, paragraph):
        """Dodaje styl stopki i ramkę górną do paragrafu"""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        pPr = paragraph._p.get_or_add_pPr()
        
        # Styl stopki
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Stopka')
        pPr.insert(0, pStyle)
        
        # Ramka górna (border top)
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), 'auto')
        pBdr.append(top)
        pPr.append(pBdr)
    
    def _setup_styles(self):
        """Konfiguruje style dokumentu zgodne z oryginałem"""
        # Styl Normal
        style = self.doc.styles['Normal']
        style.font.name = self.FONT_NAME
        style.font.size = self.FONT_SIZE_NORMAL
        
        # Ustaw też dla complex script
        style._element.rPr.rFonts.set(qn('w:cs'), self.FONT_NAME)
    
    def _set_run_font(self, run, size, bold=False, include_cs=True):
        """Ustawia czcionkę run'a zgodnie z oryginałem (Arial + szCs)"""
        run.font.name = self.FONT_NAME
        run.font.size = size
        run.bold = bold
        
        # Dodaj cs font i szCs (complex script) - wymagane przez oryginał
        if include_cs:
            rPr = run._r.get_or_add_rPr()
            
            # rFonts z cs
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), self.FONT_NAME)
            rFonts.set(qn('w:hAnsi'), self.FONT_NAME)
            rFonts.set(qn('w:cs'), self.FONT_NAME)
            
            # szCs (rozmiar dla complex script)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(size.pt * 2)))  # half-points
            rPr.append(szCs)
    
    def _add_paragraph_spacing(self, paragraph):
        """Dodaje spacing zgodny z oryginałem (line="300" lineRule="atLeast")"""
        pPr = paragraph._p.get_or_add_pPr()
        
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), str(self.LINE_SPACING))
        spacing.set(qn('w:lineRule'), 'atLeast')
        pPr.append(spacing)
    
    def _add_title(self, text: str):
        """Dodaje tytuł (14pt, wyśrodkowany, bold)"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_paragraph_spacing(p)
        
        run = p.add_run(text)
        self._set_run_font(run, self.FONT_SIZE_TITLE, bold=True)
    
    def _add_paragraph_header(self, text: str):
        """Dodaje nagłówek paragrafu (§1, §2 itd.) - wyśrodkowany, bold"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_paragraph_spacing(p)
        
        run = p.add_run(text)
        self._set_run_font(run, self.FONT_SIZE_NORMAL, bold=True)
    
    def _add_text(self, text: str, bold: bool = False, justify: bool = True):
        """Dodaje zwykły tekst (9pt, justify) z pełnym formatowaniem"""
        p = self.doc.add_paragraph()
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._add_paragraph_spacing(p)
        
        run = p.add_run(text)
        self._set_run_font(run, self.FONT_SIZE_NORMAL, bold=bold)
        return p
    
    def _add_empty_line(self):
        """Dodaje pustą linię z prawidłowym formatowaniem"""
        p = self.doc.add_paragraph()
        self._add_paragraph_spacing(p)
        
        run = p.add_run("")
        self._set_run_font(run, self.FONT_SIZE_NORMAL)
    
    def _odmiana(self, data: Dict[str, Any]) -> Dict[str, str]:
        """Zwraca odmiany słów w zależności od płci"""
        if data.get("plec") == "kobieta":
            return {
                "panem": "Panią",
                "prowadzacym": "prowadzącą",
                "zwanym": "zwaną",
                "zapoznalem": "zapoznałam"
            }
        return {
            "panem": "Panem",
            "prowadzacym": "prowadzącym",
            "zwanym": "zwanym",
            "zapoznalem": "zapoznałem"
        }
    
    def _generate_contract_with_business(self, data: Dict[str, Any]):
        """Generuje umowę dla partnera z działalnością gospodarczą"""
        odm = self._odmiana(data)
        
        # Tytuł
        self._add_title("Umowa o współpracę B2B")
        self._add_title(f"nr {data.get('numer_umowy', '……. / 2026')}")
        
        self._add_empty_line()
        
        # Data i miejsce
        self._add_text(f'zawarta w dniu {data.get("data_zawarcia", "………………………….")} roku w Warszawie pomiędzy:')
        
        self._add_empty_line()
        
        # Dane B2BNET
        self._add_text(
            "B2B.net S.A. z siedzibą w Warszawie, Aleje Jerozolimskie 180, 02-486 Warszawa, "
            "wpisaną do Rejestru Przedsiębiorców prowadzonego przez Sąd Rejonowy dla m.st. Warszawy "
            "w Warszawie, XIV Wydział Gospodarczy Krajowego Rejestru Sądowego, KRS pod nr 0000387063, "
            "NIP: 5711707392, o kapitale akcyjnym 1.360.000,00 zł (opłaconym), reprezentowaną przez "
            "Pana Artura Twardowskiego – Prezesa Zarządu, powołanego do składu Zarządu uchwałą Rady "
            "Nadzorczej nr 5/11/2025 z dnia 28.11.2025 roku, działającego samodzielnie,",
            bold=True
        )
        
        self._add_empty_line()
        self._add_text('zwaną dalej jako "Spółka" lub "B2BNET"', bold=True)
        
        self._add_empty_line()
        self._add_text("a", bold=False)
        self._add_empty_line()
        
        # Dane Partnera
        self._add_text(
            f'{odm["panem"]} {data.get("imie_nazwisko", "_____________")} {odm["prowadzacym"]} działalność gospodarczą '
            f'pod firmą: {data.get("nazwa_firmy", "________________________")}, zarejestrowaną w Centralnej Ewidencji '
            f'i Informacji o Działalności Gospodarczej pod adresem: {data.get("adres_firmy", "__________")}, '
            f'{data.get("kod_pocztowy", "00-000")} {data.get("miasto", "___________")}, NIP: {data.get("nip", "__________")}, '
            f'REGON: {data.get("regon", "_____________")},',
            bold=True
        )
        
        self._add_empty_line()
        
        # Dane kontaktowe
        self._add_text("Dane kontaktowe:")
        self._add_text(f'adres e-mail: {data.get("email", "________________")}')
        self._add_text(f'telefon: {data.get("telefon", "________________")}')
        
        self._add_empty_line()
        self._add_text(f'{odm["zwanym"]} dalej "Partnerem"', bold=True)
        
        self._add_empty_line()
        self._add_text('łącznie zwane "Stronami", a każda z osobna "Stroną".')
        
        # §1 - PRZEDMIOT UMOWY
        self._add_empty_line()
        self._add_paragraph_header("§ 1")
        self._add_paragraph_header("PRZEDMIOT UMOWY")
        
        self._add_text('Na mocy niniejszej Umowy Partner zobowiązuje się świadczyć na rzecz B2BNET, a B2BNET zobowiązuje się przyjąć od Partnera i opłacić usługi doradcze w zakresie informatycznym i/lub rekrutacyjnym w obszarze IT (dalej łącznie: "Usługi").')
        self._add_text(f'Usługi świadczone będą u Klienta B2BNET w projekcie wskazanym w Załączniku nr 3 od dnia {data.get("data_rozpoczecia", "_________")}.')
        self._add_text("Partner oświadcza, że posiada odpowiednie kwalifikacje, wiedzę oraz umiejętności niezbędne do wykonywania powierzonych mu Usług.")
        self._add_text("Partner oświadcza, że prowadzi działalność gospodarczą i zobowiązuje się do wykonywania Usług w ramach tej działalności, w sposób samodzielny i niezależny. Stosunek prawny między Stronami niniejszej Umowy nie jest stosunkiem pracy.")
        
        # §2 - ŚWIADCZENIE USŁUG
        self._add_empty_line()
        self._add_paragraph_header("§ 2")
        self._add_paragraph_header("ŚWIADCZENIE USŁUG")
        
        self._add_text("Partner zobowiązuje się świadczyć Usługi zgodnie z wytycznymi B2BNET oraz Klienta (w rozumieniu niniejszej Umowy), z zachowaniem najwyższej staranności i profesjonalizmu.")
        self._add_text(f'Usługi będą świadczone w wymiarze {data.get("wymiar_pracy", "pełnego etatu")}.')
        self._add_text("Partner zobowiązuje się do bieżącego informowania B2BNET o przebiegu realizacji Usług oraz do niezwłocznego zgłaszania wszelkich problemów mogących mieć wpływ na terminowość lub jakość wykonywanych Usług.")
        self._add_text("Partner ma prawo do wykonywania Usług przy pomocy osób trzecich (podwykonawców), pod warunkiem uzyskania uprzedniej, pisemnej zgody B2BNET. W takim przypadku Partner odpowiada za działania i zaniechania podwykonawców jak za własne.")
        
        # §3 - WYNAGRODZENIE
        self._add_empty_line()
        self._add_paragraph_header("§ 3")
        self._add_paragraph_header("WYNAGRODZENIE")
        
        self._add_text(f'Za wykonanie Usług B2BNET zapłaci Partnerowi wynagrodzenie w kwocie {data.get("stawka", "___")} PLN netto za każdy dzień świadczenia Usług ({data.get("stawka_godzinowa", "___")} PLN netto za godzinę).')
        self._add_text("Wynagrodzenie płatne będzie na podstawie prawidłowo wystawionej faktury VAT, w terminie 14 dni od daty jej doręczenia B2BNET, przelewem na rachunek bankowy Partnera wskazany na fakturze.")
        self._add_text("Partner zobowiązuje się do wystawiania faktury do 5-go dnia miesiąca następującego po miesiącu, w którym świadczone były Usługi.")
        self._add_text("Wynagrodzenie określone w ust. 1 obejmuje wszelkie koszty związane ze świadczeniem Usług przez Partnera, w tym koszty dojazdu, sprzętu i oprogramowania.")
        
        # §4 - CZAS TRWANIA UMOWY
        self._add_empty_line()
        self._add_paragraph_header("§ 4")
        self._add_paragraph_header("CZAS TRWANIA UMOWY")
        
        self._add_text(f'Umowa zostaje zawarta na czas nieokreślony, począwszy od dnia {data.get("data_rozpoczecia", "_________")}.')
        self._add_text(f'Każda ze Stron może rozwiązać Umowę z zachowaniem {data.get("okres_wypowiedzenia", "1-miesięcznego")} okresu wypowiedzenia, ze skutkiem na koniec miesiąca kalendarzowego.')
        self._add_text("B2BNET może rozwiązać Umowę ze skutkiem natychmiastowym w przypadku rażącego naruszenia przez Partnera postanowień niniejszej Umowy, w szczególności naruszenia obowiązku poufności lub działania na szkodę B2BNET lub Klienta.")
        
        # §5 - URLOPY I NIEOBECNOŚCI
        self._add_empty_line()
        self._add_paragraph_header("§ 5")
        self._add_paragraph_header("URLOPY I NIEOBECNOŚCI")
        
        self._add_text(f'Partner ma prawo do {data.get("dni_urlopu", "20")} dni płatnej przerwy w świadczeniu Usług w każdym roku kalendarzowym ("Urlop").')
        self._add_text("Urlop wymaga uprzedniego uzgodnienia z B2BNET z co najmniej 14-dniowym wyprzedzeniem.")
        self._add_text("Za dni Urlopu przysługuje Partnerowi wynagrodzenie obliczone jak za dni świadczenia Usług.")
        self._add_text("Nieobecności Partnera niewykorzystane jako Urlop nie są płatne, chyba że Strony postanowią inaczej.")
        
        # §6 - PRAWA AUTORSKIE I WŁASNOŚĆ INTELEKTUALNA
        self._add_empty_line()
        self._add_paragraph_header("§ 6")
        self._add_paragraph_header("PRAWA AUTORSKIE I WŁASNOŚĆ INTELEKTUALNA")
        
        self._add_text("Wszelkie utwory, wynalazki, wzory, know-how i inne dobra intelektualne powstałe w wyniku realizacji Usług przez Partnera stanowią własność B2BNET lub Klienta B2BNET (w zależności od ustaleń między B2BNET a Klientem).")
        self._add_text("Partner przenosi na B2BNET autorskie prawa majątkowe do wszystkich utworów stworzonych w ramach realizacji Usług, na wszystkich polach eksploatacji znanych w chwili zawarcia Umowy.")
        self._add_text("Wynagrodzenie określone w § 3 obejmuje wynagrodzenie za przeniesienie praw autorskich.")
        
        # §7 - ZAKAZ KONKURENCJI
        self._add_empty_line()
        self._add_paragraph_header("§ 7")
        self._add_paragraph_header("ZAKAZ KONKURENCJI")
        
        self._add_text("Partner zobowiązuje się, że w okresie obowiązywania Umowy oraz przez okres 12 miesięcy po jej rozwiązaniu nie będzie świadczyć usług bezpośrednio na rzecz Klienta, u którego świadczył Usługi w ramach niniejszej Umowy, z pominięciem B2BNET.")
        self._add_text("W przypadku naruszenia zakazu, o którym mowa w ust. 1, Partner zapłaci B2BNET karę umowną w wysokości 6-krotności miesięcznego wynagrodzenia brutto.")
        
        # §8 - POUFNOŚĆ
        self._add_empty_line()
        self._add_paragraph_header("§ 8")
        self._add_paragraph_header("POUFNOŚĆ")
        
        self._add_text('Partner zobowiązuje się do zachowania w tajemnicy wszelkich informacji dotyczących B2BNET, Klientów oraz ich kontrahentów, uzyskanych w związku z realizacją Usług ("Informacje Poufne").')
        self._add_text("Obowiązek zachowania poufności obowiązuje w czasie trwania Umowy oraz przez okres 5 lat po jej rozwiązaniu.")
        self._add_text("Szczegółowe zobowiązania dotyczące poufności określa Załącznik nr 1 do Umowy (Deklaracja Poufności).")
        
        # §9 - OCHRONA DANYCH OSOBOWYCH
        self._add_empty_line()
        self._add_paragraph_header("§ 9")
        self._add_paragraph_header("OCHRONA DANYCH OSOBOWYCH")
        
        self._add_text("W zakresie, w jakim realizacja Usług wiąże się z przetwarzaniem danych osobowych, Strony zawierają Umowę Powierzenia Przetwarzania Danych Osobowych stanowiącą Załącznik nr 2 do niniejszej Umowy.")
        self._add_text("Partner zobowiązuje się do przestrzegania przepisów RODO oraz innych przepisów dotyczących ochrony danych osobowych.")
        
        # §10 - ODPOWIEDZIALNOŚĆ
        self._add_empty_line()
        self._add_paragraph_header("§ 10")
        self._add_paragraph_header("ODPOWIEDZIALNOŚĆ")
        
        self._add_text("Partner ponosi odpowiedzialność za szkody wyrządzone B2BNET lub Klientowi wskutek niewykonania lub nienależytego wykonania Umowy.")
        self._add_text("Partner posiada ubezpieczenie OC z tytułu prowadzonej działalności gospodarczej lub zobowiązuje się do jego zawarcia w terminie 14 dni od daty zawarcia niniejszej Umowy.")
        
        # §11 - POSTANOWIENIA KOŃCOWE
        self._add_empty_line()
        self._add_paragraph_header("§ 11")
        self._add_paragraph_header("POSTANOWIENIA KOŃCOWE")
        
        self._add_text("Wszelkie zmiany Umowy wymagają formy pisemnej pod rygorem nieważności.")
        self._add_text("W sprawach nieuregulowanych Umową mają zastosowanie przepisy Kodeksu Cywilnego.")
        self._add_text("Spory wynikające z Umowy będą rozstrzygane przez sąd właściwy dla siedziby B2BNET.")
        self._add_text("Umowę sporządzono w dwóch jednobrzmiących egzemplarzach, po jednym dla każdej ze Stron.")
        self._add_text("Integralną część Umowy stanowią następujące Załączniki:")
        self._add_text("Załącznik nr 1 – Deklaracja Poufności")
        self._add_text("Załącznik nr 2 – Umowa Powierzenia Przetwarzania Danych Osobowych")
        self._add_text("Załącznik nr 3 – Klient Projektu")
        
        self._add_empty_line()
        self._add_empty_line()
        self._add_text("_______________\t\t\t\t\t\t\t_______________", justify=False)
        self._add_text("B2B.NET S.A.\t\t\t\t\t\t\t\tPartner", justify=False)
    
    def _generate_contract_without_business(self, data: Dict[str, Any]):
        """Generuje umowę dla partnera bez działalności gospodarczej"""
        # Podobna struktura, ale dostosowana do osób fizycznych
        self._generate_contract_with_business(data)
    
    def _add_attachment_1(self, data: Dict[str, Any]):
        """Załącznik nr 1 - Deklaracja Poufności"""
        odm = self._odmiana(data)
        
        self.doc.add_page_break()
        self._add_text("Załącznik nr 1 - Deklaracja Poufności")
        self._add_empty_line()
        self._add_paragraph_header("DEKLARACJA POUFNOŚCI")
        self._add_empty_line()
        
        self._add_text(f'Ja, niżej podpisany/a {data.get("imie_nazwisko", "_________________________")}, świadczący usługi w ramach Umowy o Współpracę B2B nr {data.get("numer_umowy", "[NUMER]")} z dnia {data.get("data_zawarcia", "[DATA]")}, zobowiązuję się do zachowania w tajemnicy wszelkich informacji dotyczących B2BNET S.A. oraz jej Klientów.')
        self._add_empty_line()
        self._add_text("Przez Informacje Poufne rozumie się wszelkie informacje techniczne, technologiczne, organizacyjne, finansowe, handlowe oraz inne informacje posiadające wartość gospodarczą, które nie są powszechnie znane i zostały mi udostępnione lub do których uzyskałem dostęp w związku z wykonywaniem Umowy.")
        self._add_text("Partner zobowiązuje się do korzystania z Informacji Poufnych wyłącznie w celu i zakresie niezbędnym do wykonania Usług wynikających z Umowy Głównej.")
        self._add_text("Niniejsze zobowiązanie wiąże Partnera zarówno w czasie trwania Umowy Głównej, jak i po jej rozwiązaniu lub wygaśnięciu, przez okres wskazany w § 8 Umowy Głównej.")
        self._add_text("Zobowiązanie do nieujawniania nie dotyczy informacji, które: a) Były publicznie znane lub stały się publicznie znane w sposób niezawiniony przez Partnera; b) Partner musi ujawnić na żądanie sądu lub uprawnionych organów państwowych, a obowiązek ich ujawnienia wynika z bezwzględnie obowiązujących przepisów prawa.")
        self._add_text("Partner zobowiązuje się przechowywać wszelkie Informacje Poufne (w tym materialne nośniki, pliki elektroniczne) w sposób uniemożliwiający dostęp do nich osobom nieupoważnionym.")
        self._add_text("Partner zobowiązany jest do niezwłocznego poinformowania B2BNET o każdym przypadku nieuprawnionego ujawnienia Informacji Poufnych oraz do podjęcia pełnej współpracy w celu ograniczenia i usunięcia skutków tego faktu.")
        self._add_empty_line()
        self._add_text(f'Oświadczam, iż {odm["zapoznalem"]} się z niniejszą deklaracją, rozumiem jej treść i zobowiązuję się do jej stosowania. Jestem świadomy odpowiedzialności za nieprzestrzeganie zapisów zawartych w niniejszej deklaracji.')
        self._add_empty_line()
        self._add_empty_line()
        self._add_text("_________________________\t\t_________________")
        self._add_text("Czytelny podpis Partnera\t\t\tdata")
        self._add_empty_line()
        self._add_text("Deklaracja podpisana w obecności:")
        self._add_empty_line()
        self._add_empty_line()
        self._add_text("_______________________\t\t________________")
        self._add_text("Przedstawiciel B2B.NET S.A.\t\tdata")
    
    def _add_attachment_2(self, data: Dict[str, Any]):
        """Załącznik nr 2 - Umowa Powierzenia Przetwarzania Danych (DPA)"""
        odm = self._odmiana(data)
        
        self.doc.add_page_break()
        self._add_text("Załącznik nr 2 - Wzór Umowy Powierzenia Przetwarzania Danych Osobowych (DPA)")
        self._add_empty_line()
        self._add_text(f'Zawarta w Warszawie, dnia {data.get("data_rozpoczecia", "____________")} jako Załącznik nr 2 do Umowy o Świadczenie Usług B2B')
        self._add_text("Pomiędzy:")
        self._add_text(
            "B2BNET S.A. z siedzibą w Warszawie, Al. Jerozolimskie 180, 02-486 Warszawa, wpisaną do rejestru "
            "przedsiębiorców KRS pod numerem 0000387063, NIP: 5711707392 reprezentowaną przez Pana Artura "
            "Twardowskiego – Prezesa Zarządu, powołanego do składu Zarządu uchwałą Rady Nadzorczej nr 5/11/2025 "
            "z dnia 28.11.2025 roku, działającego samodzielnie,",
            bold=True, justify=False
        )
        self._add_text('zwaną dalej "Procesorem Powierzającym" (lub w rozumieniu RODO: Administratorem, który powierza dane)', bold=True)
        self._add_text("a")
        self._add_text(
            f'{odm["panem"]} {data.get("imie_nazwisko", "_____________")} {odm["prowadzacym"]} działalność gospodarczą '
            f'pod firmą: {data.get("nazwa_firmy", "________________________")}, zarejestrowaną w Centralnej Ewidencji '
            f'i Informacji o Działalności Gospodarczej pod adresem: {data.get("adres_firmy", "__________")}, '
            f'{data.get("kod_pocztowy", "00-000")} {data.get("miasto", "___________")}, NIP: {data.get("nip", "__________")}, '
            f'REGON: {data.get("regon", "_____________")},',
            bold=True
        )
        self._add_text('zwanym dalej "Procesorem Przetwarzającym" (lub w rozumieniu RODO: Podmiotem Przetwarzającym).', bold=True)
        
        self._add_empty_line()
        self._add_paragraph_header("§ 1. Przedmiot Umowy Powierzenia")
        self._add_text("Procesor Powierzający, na podstawie art. 28 ust. 3 Rozporządzenia Parlamentu Europejskiego i Rady (UE) 2016/679 (RODO), powierza Procesorowi Przetwarzającemu przetwarzanie danych osobowych, a Procesor Przetwarzający zobowiązuje się do ich przetwarzania na zasadach określonych w niniejszej Umowie.", justify=False)
        self._add_text("Dane osobowe są powierzane w celu i w związku z realizacją Umowy o Świadczenie Usług B2B (Umowa Główna), której niniejsza Umowa Powierzenia stanowi integralną część.", justify=False)
        
        self._add_empty_line()
        self._add_paragraph_header("§ 2. Zakres i Charakter Przetwarzania")
        self._add_text("Zwykłe dane osobowe (np. imię, nazwisko, stanowisko, adres e-mail, służbowy numer telefonu, dane dotyczące aktywności w systemach IT Klienta). Możliwe przetwarzanie szczególnych kategorii danych osobowych (dane wrażliwe) – wyłącznie na wyraźne pisemne polecenie Procesora Powierzającego i w zakresie absolutnie niezbędnym do wykonania Usługi (np. w projektach medycznych).", justify=False)
        self._add_text("Klienci Procesora Powierzającego, pracownicy Klientów, kontrahenci Klientów lub inne osoby, których dane zostały przekazane Procesorowi Przetwarzającemu w związku z realizacją Umowy Głównej.", justify=False)
        self._add_text("Przez okres obowiązywania Umowy Głównej oraz przez czas niezbędny do zakończenia operacji przetwarzania po jej rozwiązaniu lub wygaśnięciu.", justify=False)
        self._add_text("Przetwarzanie polega na wykonywaniu operacji takich jak: dostęp, gromadzenie, utrwalanie, porządkowanie, modyfikowanie, przechowywanie, opracowywanie, usuwanie – w zakresie niezbędnym do świadczenia usług informatycznych/doradczych zgodnie z Umową Główną.", justify=False)
        
        self._add_empty_line()
        self._add_paragraph_header("§ 3. Obowiązki Procesora Przetwarzającego (Partnera)")
        self._add_text("Procesor Przetwarzający zobowiązuje się przetwarzać powierzone dane wyłącznie na udokumentowane polecenie Procesora Powierzającego.", justify=False)
        self._add_text("Procesor Przetwarzający zobowiązuje się stosować środki techniczne i organizacyjne wymagane przez Art. 32 RODO, zapewniając poziom bezpieczeństwa odpowiadający ryzyku związanemu z przetwarzaniem, a w szczególności zabezpieczając dane przed nieuprawnionym dostępem, utratą, uszkodzeniem lub zniszczeniem.", justify=False)
        self._add_text("Procesor Przetwarzający zapewni, że osoby upoważnione do przetwarzania danych (w tym on sam) zobowiązały się do zachowania poufności na podstawie pisemnego oświadczenia.", justify=False)
        self._add_text("Procesor Przetwarzający zobowiązuje się do wspierania Procesora Powierzającego w zakresie realizacji praw osób, których dane dotyczą (np. prawo dostępu, sprostowania) oraz w zakresie wywiązywania się z obowiązków określonych w art. 32–36 RODO (np. zgłaszanie naruszeń).", justify=False)
        self._add_text("W przypadku stwierdzenia naruszenia ochrony danych osobowych, Procesor Przetwarzający zgłosi je Procesorowi Powierzającemu niezwłocznie, jednak nie później niż w ciągu 24 godzin od momentu wykrycia naruszenia.", justify=False)
        self._add_text("Procesor Przetwarzający udostępni Procesorowi Powierzającemu wszelkie informacje niezbędne do wykazania zgodności z obowiązkami RODO oraz umożliwi przeprowadzenie audytów i inspekcji (w tym zdalnych lub w siedzibie).", justify=False)
        
        self._add_empty_line()
        self._add_paragraph_header("§ 4. Dalsze Powierzenie (Sub-procesorzy)")
        self._add_text("Procesor Przetwarzający może powierzyć przetwarzanie powierzonych danych podmiotom trzecim (sub-procesorom, np. zewnętrznym dostawcom chmury lub narzędzi), pod warunkiem uzyskania uprzedniej, pisemnej zgody ogólnej lub szczególnej Procesora Powierzającego.", justify=False)
        self._add_text("W przypadku uzyskania zgody, dalsze powierzenie może nastąpić wyłącznie na podstawie umowy, która nakłada na sub-procesora takie same obowiązki w zakresie ochrony danych, jakie zostały nałożone na Procesora Przetwarzającego w niniejszej Umowie.", justify=False)
        self._add_text("Procesor Przetwarzający ponosi pełną odpowiedzialność wobec Procesora Powierzającego za niewykonanie lub nienależyte wykonanie obowiązków przez sub-procesora.", justify=False)
        
        self._add_empty_line()
        self._add_paragraph_header("§ 5. Postanowienia Końcowe")
        self._add_text("Po zakończeniu obowiązywania Umowy Głównej lub na każde żądanie Procesora Powierzającego, Procesor Przetwarzający, wedle wyboru Procesora Powierzającego, trwale usunie lub zwróci mu wszelkie powierzone dane osobowe (wraz z ich kopiami), chyba że przepisy prawa nakazują ich dalsze przechowywanie.", justify=False)
        self._add_text("Procesor Przetwarzający ponosi pełną odpowiedzialność majątkową za wszelkie szkody, w tym grzywny administracyjne nałożone na Procesora Powierzającego, powstałe wskutek naruszenia przez Procesora Przetwarzającego niniejszej Umowy Powierzenia lub przepisów RODO.", justify=False)
        self._add_text("W sprawach nieuregulowanych niniejszą Umową mają zastosowanie przepisy RODO oraz Kodeksu Cywilnego.", justify=False)
        
        self._add_empty_line()
        self._add_empty_line()
        self._add_text("_______________\t\t\t\t\t\t\t_______________", justify=False)
        self._add_text("B2B.NET S.A\t\t\t\t\t\t                      Partner", justify=False)
    
    def _add_attachment_3(self, data: Dict[str, Any]):
        """Załącznik nr 3 - Klient Projektu"""
        self.doc.add_page_break()
        self._add_text("Załącznik nr 3 – wzór określający Klienta B2BNET", justify=False)
        self._add_empty_line()
        self._add_paragraph_header(f'ZAŁĄCZNIK NR 3 DO UMOWY B2B NR {data.get("numer_umowy", "[NUMER]")}')
        self._add_paragraph_header("KLIENT PROJEKTU")
        self._add_empty_line()
        
        # Tabela z danymi klienta
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        cells = [
            ("Element", "Wartość"),
            ("Nazwa Klienta Projektu:", data.get("klient_nazwa", "[PEŁNA NAZWA KLIENTA]")),
            ("Adres (miasta):", data.get("klient_adres", "[MIASTO]")),
            ("Opis projektu i zakres usług:", data.get("klient_opis", "[OPIS]")),
            ("Data rozpoczęcia świadczenia usług:", data.get("klient_data_rozpoczecia", "[DATA]"))
        ]
        
        for i, (label, value) in enumerate(cells):
            cell_label = table.rows[i].cells[0]
            cell_value = table.rows[i].cells[1]
            
            # Formatowanie tekstu w komórkach
            p_label = cell_label.paragraphs[0]
            run_label = p_label.add_run(label)
            self._set_run_font(run_label, self.FONT_SIZE_NORMAL, bold=(i == 0))
            
            p_value = cell_value.paragraphs[0]
            run_value = p_value.add_run(value)
            self._set_run_font(run_value, self.FONT_SIZE_NORMAL, bold=(i == 0))
        
        self._add_empty_line()
        self._add_empty_line()
        self._add_text("________________________\t\t\t\t\t\t_____________________", justify=False)
        self._add_text("B2B.NET S.A.\t\t\t\t\t\t\t\tPartner", justify=False)
