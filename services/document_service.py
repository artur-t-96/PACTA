"""
Document Service - handles template storage, field detection, and document generation.
Uses docxtpl for template rendering with Jinja2-style placeholders.
"""
import os
import re
import uuid
import logging
from io import BytesIO
from typing import List, Tuple, Optional, Dict, Any

from docxtpl import DocxTemplate
from docx import Document

import config

logger = logging.getLogger(__name__)

# Regex pattern for detecting placeholders: {{FIELD_NAME}}
PLACEHOLDER_PATTERN = re.compile(r'\{\{([A-Z0-9_]+)\}\}')


def ensure_storage_dirs():
    """Create storage directories if they don't exist."""
    for dir_path in [config.TEMPLATES_DIR, config.GENERATED_DIR]:
        try:
            os.makedirs(dir_path, exist_ok=True)
            logger.info(f"Storage directory ensured: {dir_path}")
        except PermissionError:
            logger.warning(f"Cannot create directory {dir_path}, using fallback")
            # Fallback to local directories for development
            fallback = dir_path.replace("/var/data", "./data")
            os.makedirs(fallback, exist_ok=True)


def get_templates_dir() -> str:
    """Get the templates directory, creating fallback if needed."""
    if os.path.exists(config.TEMPLATES_DIR) and os.access(config.TEMPLATES_DIR, os.W_OK):
        return config.TEMPLATES_DIR
    # Fallback for local development
    fallback = config.TEMPLATES_DIR.replace("/var/data", "./data")
    os.makedirs(fallback, exist_ok=True)
    return fallback


def get_generated_dir() -> str:
    """Get the generated documents directory, creating fallback if needed."""
    if os.path.exists(config.GENERATED_DIR) and os.access(config.GENERATED_DIR, os.W_OK):
        return config.GENERATED_DIR
    # Fallback for local development
    fallback = config.GENERATED_DIR.replace("/var/data", "./data")
    os.makedirs(fallback, exist_ok=True)
    return fallback


def extract_fields_from_docx(file_path: str) -> List[str]:
    """
    Extract placeholder fields from a DOCX document.
    Looks for patterns like {{FIELD_NAME}} in the document text.

    Args:
        file_path: Path to the DOCX file

    Returns:
        List of unique field names found in the document
    """
    fields = set()

    try:
        doc = Document(file_path)

        # Extract from paragraphs
        for para in doc.paragraphs:
            matches = PLACEHOLDER_PATTERN.findall(para.text)
            fields.update(matches)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = PLACEHOLDER_PATTERN.findall(para.text)
                        fields.update(matches)

        # Extract from headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        matches = PLACEHOLDER_PATTERN.findall(para.text)
                        fields.update(matches)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        matches = PLACEHOLDER_PATTERN.findall(para.text)
                        fields.update(matches)

        logger.info(f"Extracted {len(fields)} fields from {file_path}: {fields}")
        return sorted(list(fields))

    except Exception as e:
        logger.error(f"Error extracting fields from {file_path}: {e}")
        raise ValueError(f"Nie udalo sie odczytac szablonu: {str(e)}")


def save_template_file(file_content: bytes, original_filename: str) -> Tuple[str, List[str]]:
    """
    Save a template file to disk and extract its fields.

    Args:
        file_content: Binary content of the uploaded file
        original_filename: Original filename from upload

    Returns:
        Tuple of (file_path, list of detected fields)
    """
    # Generate unique filename
    file_id = str(uuid.uuid4())[:8]
    safe_name = re.sub(r'[^a-zA-Z0-9._-]', '_', original_filename)
    filename = f"{file_id}_{safe_name}"

    templates_dir = get_templates_dir()
    file_path = os.path.join(templates_dir, filename)

    # Save file to disk
    with open(file_path, 'wb') as f:
        f.write(file_content)

    logger.info(f"Template saved to: {file_path}")

    # Extract fields
    fields = extract_fields_from_docx(file_path)

    return file_path, fields


def generate_document_from_template(template_path: str, data: Dict[str, Any]) -> bytes:
    """
    Generate a document from a template with substituted data.

    Args:
        template_path: Path to the template file
        data: Dictionary of field values to substitute

    Returns:
        Generated document as bytes
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Szablon nie istnieje: {template_path}")

    try:
        # Load template
        doc = DocxTemplate(template_path)

        # Render with data (docxtpl uses Jinja2 syntax)
        # Convert {{FIELD}} to Jinja2 {{ FIELD }} format internally
        doc.render(data)

        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        logger.info(f"Document generated from template: {template_path}")
        return output.getvalue()

    except Exception as e:
        logger.error(f"Error generating document from {template_path}: {e}")
        raise ValueError(f"Blad generowania dokumentu: {str(e)}")


def delete_template_file(file_path: str) -> bool:
    """
    Delete a template file from disk.

    Args:
        file_path: Path to the file to delete

    Returns:
        True if deleted, False if file didn't exist
    """
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
            logger.info(f"Template deleted: {file_path}")
            return True
        except Exception as e:
            logger.error(f"Error deleting template {file_path}: {e}")
            raise
    return False


def convert_placeholders_to_jinja(template_path: str) -> str:
    """
    Convert {{FIELD}} placeholders to Jinja2 {{ FIELD }} format.
    Creates a temporary converted file and returns its path.

    Note: docxtpl already supports Jinja2 syntax, so this may not be needed
    if templates already use {{ field }} format.
    """
    # For now, assume templates use {{ field }} format compatible with docxtpl
    return template_path


def validate_template_fields(template_fields: List[str], provided_data: Dict[str, Any]) -> List[str]:
    """
    Validate that all required template fields are provided.

    Args:
        template_fields: List of fields required by template
        provided_data: Dictionary of provided field values

    Returns:
        List of missing field names
    """
    missing = []
    for field in template_fields:
        if field not in provided_data or not provided_data[field]:
            missing.append(field)
    return missing


def format_field_label(field_name: str) -> str:
    """
    Convert field name to human-readable label.
    Example: IMIE_NAZWISKO -> Imie nazwisko

    Args:
        field_name: Field name in UPPER_SNAKE_CASE

    Returns:
        Human-readable label
    """
    return field_name.replace("_", " ").capitalize()
