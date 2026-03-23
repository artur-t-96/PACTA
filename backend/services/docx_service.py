"""
DOCX Service — python-docx based contract generation
"""
import os
import re
import copy
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt
from typing import Optional


TEMPLATE_PATH = os.getenv(
    "TEMPLATE_PATH",
    "/var/data/templates/Umowa_B2B_draft_2026_od_02.03.2026.docx",
)
OUTPUT_DIR = os.getenv("OUTPUT_DIR", "/var/data/output/contracts")


def _format_date_pl(date_str: str) -> str:
    """Convert YYYY-MM-DD or DD.MM.YYYY to Polish formatted date."""
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%d-%m-%Y"):
        try:
            dt = datetime.strptime(date_str, fmt)
            months = [
                "stycznia", "lutego", "marca", "kwietnia", "maja", "czerwca",
                "lipca", "sierpnia", "września", "października", "listopada", "grudnia"
            ]
            return f"{dt.day} {months[dt.month - 1]} {dt.year}"
        except ValueError:
            continue
    return date_str


def _format_date_short(date_str: str) -> str:
    """Convert to DD.MM.YYYY format."""
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%d-%m-%Y"):
        try:
            dt = datetime.strptime(date_str, fmt)
            return dt.strftime("%d.%m.%Y")
        except ValueError:
            continue
    return date_str


def _replace_in_paragraph(paragraph, replacements: dict):
    """Replace placeholders in a paragraph preserving formatting.
    Handles cross-run replacements by joining all runs, replacing, then redistributing."""
    full_text = paragraph.text
    needs_replace = any(k in full_text for k in replacements)
    
    if not needs_replace:
        return
    
    # First try simple run-level replacement
    for run in paragraph.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, str(new))
    
    # If still has placeholders, do cross-run replacement
    full_text_after = paragraph.text
    still_needs = any(k in full_text_after for k in replacements)
    
    if still_needs and paragraph.runs:
        # Join all run text, replace, put in first run, clear rest
        joined = "".join(run.text for run in paragraph.runs)
        for old, new in replacements.items():
            joined = joined.replace(old, str(new))
        
        # Keep formatting of first run, clear others
        first_run = paragraph.runs[0]
        first_run.text = joined
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_table(table, replacements: dict):
    """Replace placeholders in table cells."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_in_paragraph(para, replacements)


def generate_contract_docx(
    contract_number: str,
    imie: str,
    nazwisko: str,
    firma: str,
    nip: str,
    regon: str,
    adres: str,
    email: str,
    tel: str,
    rola: str,
    stawka: float,
    klient: str,
    data_startu: str,
    it_area: str,
    opis_projektu: str,
    miasto_klienta: str,
) -> str:
    """
    Fill in the B2B contract template and save to OUTPUT_DIR.
    Returns path to generated file.
    """
    doc = Document(TEMPLATE_PATH)

    # Today's date
    today = datetime.now()
    today_str = _format_date_pl(today.strftime("%Y-%m-%d"))
    today_dot = today.strftime("%d.%m.%Y")
    start_date_pl = _format_date_pl(data_startu)
    start_date_short = _format_date_short(data_startu)

    # Replacements mapping — handles all the placeholders found in the template
    contractor_full = f"{imie} {nazwisko}"
    contractor_intro = (
        f"{imie} {nazwisko} prowadzącym/cą działalność gospodarczą pod firmą: "
        f"{firma}, zarejestrowaną w Centralnej Ewidencji i Informacji o Działalności "
        f"Gospodarczej pod adresem: {adres}, NIP: {nip}"
        + (f", REGON: {regon}" if regon else "")
    )

    replacements = {
        # Contract number
        "……. / 2026": contract_number,
        "nr ……. / 2026": f"nr {contract_number}",
        "[NUMER]": contract_number,
        # Date of signing - exact placeholder from template
        "………………………….": today_dot,
        "………………………………….": today_dot,
        "……………………………………": today_dot,
        "…………………………": today_dot,
        # Contractor identification line
        "Panem/ią _____________ prowadzącym/cą działalność gospodarczą pod firmą: ________________________, zarejestrowaną w Centralnej Ewidencji i Informacji o Działalności Gospodarczej pod adresem: _________": contractor_intro,
        # Contact
        "Tel. ………. Adres e-mail: …….": f"Tel. {tel} Adres e-mail: {email}",
        # IT area in §1
        "______________(należy wpisać konkretny obszar działalności Par": f"{it_area} (specjalizacja: {rola}",
        # Rate in §6
        "……………. z": f"{stawka:.2f} z",
        # Start date in §12
        "Partner zobowiązany jest do podjęcia świadczenia Usług z dniem ………………..": f"Partner zobowiązany jest do podjęcia świadczenia Usług z dniem {start_date_pl}.",
        # DPA attachment contractor line
        "Panem/ią _____________ prowadzącym/cą działalność gospodarczą pod firmą: ________________________, zarejestrowaną w Centralnej Ewidencji i Informacji o Działalności Gospodarczej pod adresem: ________": contractor_intro,
    }

    # Process all paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Process all tables
    for table in doc.tables:
        _replace_in_table(table, replacements)

    # Fill Attachment 3 (Klient Projektu) table
    att3_replacements = {
        "[PEŁNA NAZWA KLIENTA]": klient,
        "Wpisujemy nazwy miast, które Klient wskazuje, jeżeli podany jest adres szczegółowy to również wpisujemy": miasto_klienta,
        "Wpisujemy te informacje, które są znane z procesu rekrutacyjnego": opis_projektu or f"Usługi w obszarze {it_area} — {rola}",
        "……": start_date_short,
    }
    for table in doc.tables:
        _replace_in_table(table, att3_replacements)

    # Fix Deklaracja Poufności and DPA date references
    decl_replacements = {
        "nr ……. z dnia ………": f"nr {contract_number} z dnia {today_dot}",
        # DPA date
        "dnia ____________ jako Załącznik nr 2 do Umowy o współpracę B2B nr _______": f"dnia {today_dot} jako Załącznik nr 2 do Umowy o współpracę B2B nr {contract_number}",
    }
    for para in doc.paragraphs:
        _replace_in_paragraph(para, decl_replacements)
    
    # Universal cleanup: signature lines — replace underscore patterns with names
    for para in doc.paragraphs:
        text = para.text.strip()
        # Match patterns like "___...\t\t\t___..." (two signature fields separated by tabs)
        if re.match(r'^_{5,}\s*\t+\s*_{5,}\s*$', text):
            # Pure signature line — replace with names
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = f"{contractor_full}\t\t\t\t\t\tB2B.net S.A."
        # Fix partially replaced signatures with surrounding underscores
        elif re.search(r'_{3,}[^_]+_{3,}', text) and contractor_full in text:
            cleaned = re.sub(r'_{3,}', '', text)
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = cleaned.strip()

    # Save
    safe_number = contract_number.replace("/", "_")
    safe_name = f"{nazwisko}_{imie}".replace(" ", "_")
    filename = f"Umowa_B2B_{safe_number}_{safe_name}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(output_path)

    return output_path


def get_paragraph_text(file_path: str, paragraf_ref: str) -> Optional[str]:
    """
    Find and return text of a specific paragraph by reference (e.g. "§10 ust.1").
    """
    doc = Document(file_path)
    
    # Normalize search
    para_norm = paragraf_ref.replace("§", "§").strip()
    
    # Strategy: find paragraph containing § reference
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    
    # Try to find the section header and return the next non-empty paragraph(s)
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        # Match § numbers
        if "§ 10" in text or "§10" in text or "§ 10." in text:
            if "10" in paragraf_ref:
                # Find ust.1 (first list item after §10)
                for j in range(i + 1, min(i + 15, len(paragraphs))):
                    if paragraphs[j].text.strip():
                        return paragraphs[j].text.strip()
    
    # Fallback: search by content similarity
    for para in paragraphs:
        if any(term in para.text for term in ["miesiąc", "miesięcy", "zakaz", "konkurencj"]):
            return para.text.strip()
    
    return None


def _create_version(file_path: str) -> str:
    """Create a versioned copy of the file before modification."""
    if not os.path.exists(file_path):
        return file_path
    
    versions_dir = os.path.join(os.path.dirname(file_path), "versions")
    os.makedirs(versions_dir, exist_ok=True)
    
    base = os.path.basename(file_path).replace(".docx", "")
    import glob
    existing = glob.glob(os.path.join(versions_dir, f"{base}_v*.docx"))
    version = len(existing) + 1
    
    version_path = os.path.join(versions_dir, f"{base}_v{version}.docx")
    import shutil
    shutil.copy2(file_path, version_path)
    return version_path


def modify_paragraph_in_docx(file_path: str, paragraf_ref: str, new_text: str) -> str:
    """
    Modify a paragraph in the DOCX and save a new version.
    Creates a versioned backup before modifying.
    Returns path to new file.
    """
    # Create version backup before modifying
    _create_version(file_path)
    
    doc = Document(file_path)
    
    # Find and replace the paragraph
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    modified = False
    
    para_num = re.search(r'§\s*(\d+)', paragraf_ref)
    ust_num = re.search(r'ust\.?\s*(\d+)', paragraf_ref)
    
    target_section = para_num.group(1) if para_num else None
    target_ust = ust_num.group(1) if ust_num else "1"
    
    in_target_section = False
    ust_counter = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Detect section header
        if target_section and re.search(rf'§\s*{target_section}[^\d]|§\s*{target_section}$', text):
            in_target_section = True
            ust_counter = 0
            continue
        
        # Detect next section (reset)
        if in_target_section and re.search(r'§\s*\d+', text) and not re.search(rf'§\s*{target_section}', text):
            in_target_section = False
        
        if in_target_section and text:
            ust_counter += 1
            if str(ust_counter) == target_ust:
                # Replace text in this paragraph
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)
                modified = True
                break
    
    if not modified:
        # Try table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = new_text
                            else:
                                para.add_run(new_text)
                            modified = True
                            break

    # Save as new version
    base, ext = os.path.splitext(file_path)
    new_path = f"{base}_v{int(datetime.now().timestamp())}{ext}"
    doc.save(new_path)
    return new_path
