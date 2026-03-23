"""
Template service — manage contract templates
"""
import os
import glob
from docx import Document


TEMPLATE_DIR = os.path.expanduser("~/clawd/data/legal/marta_templates")


def list_templates() -> list:
    """List available contract templates."""
    templates = []
    
    for f in glob.glob(os.path.join(TEMPLATE_DIR, "*.docx")):
        doc = Document(f)
        # Get first paragraph as title
        title = ""
        for p in doc.paragraphs[:5]:
            if p.text.strip():
                title = p.text.strip()[:100]
                break
        
        stat = os.stat(f)
        templates.append({
            "filename": os.path.basename(f),
            "path": f,
            "title": title,
            "size_kb": round(stat.st_size / 1024, 1),
            "paragraphs": sum(1 for p in doc.paragraphs if p.text.strip()),
            "tables": len(doc.tables),
        })
    
    return templates


def get_template_structure(filename: str) -> dict:
    """Get detailed structure of a template (paragraphs with §)."""
    path = os.path.join(TEMPLATE_DIR, filename)
    if not os.path.exists(path):
        return {"error": "Template not found"}
    
    doc = Document(path)
    
    sections = []
    current_section = None
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        
        if text.startswith("§") or "ZAŁĄCZNIK" in text.upper() or "UMOWA" in text.upper():
            current_section = {"header": text[:150], "content": []}
            sections.append(current_section)
        elif current_section:
            current_section["content"].append(text[:200])
    
    return {
        "filename": filename,
        "sections": sections[:30],
        "total_paragraphs": sum(1 for p in doc.paragraphs if p.text.strip()),
    }
