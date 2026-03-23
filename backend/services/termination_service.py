"""
Termination service — generate mutual termination agreements
(Porozumienie o rozwiązaniu umowy za zgodą stron)
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.getenv("OUTPUT_DIR", "./output/contracts/terminations")


def generate_termination(
    contract_number: str,
    contractor_name: str,
    contractor_company: str,
    contractor_nip: str,
    termination_date: str,
    reason: str = "",
    settlement_notes: str = "",
) -> str:
    """Generate mutual termination agreement DOCX."""
    
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    today = datetime.now().strftime("%d.%m.%Y")
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("POROZUMIENIE O ROZWIAZANIU UMOWY")
    run.bold = True
    run.font.size = Pt(14)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"do Umowy nr {contract_number}")
    run.font.size = Pt(11)
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"Warszawa, dnia {today}")
    
    doc.add_paragraph()
    doc.add_paragraph("Zawarte pomiedzzy:")
    
    p = doc.add_paragraph()
    p.add_run("B2BNET S.A.").bold = True
    p.add_run(" z siedziba w Warszawie, Al. Jerozolimskie 180, 02-486 Warszawa, "
              "KRS: 0000387063, NIP: 5711707392, "
              "reprezentowana przez Artura Twardowskiego - Prezesa Zarzadu")
    
    doc.add_paragraph("a")
    
    p = doc.add_paragraph()
    p.add_run(contractor_name).bold = True
    nip = f", NIP: {contractor_nip}" if contractor_nip else ""
    p.add_run(f" prowadzacym/a dzialalnosc gospodarcza pod firma: {contractor_company}{nip}")
    
    doc.add_paragraph()
    
    # §1
    p = doc.add_paragraph()
    p.add_run("Par. 1").bold = True
    doc.add_paragraph(
        f"Strony zgodnie postanawiaja rozwiazac Umowe o wspolprace B2B nr {contract_number} "
        f"z dniem {termination_date}."
    )
    
    doc.add_paragraph()
    
    # §2
    p = doc.add_paragraph()
    p.add_run("Par. 2").bold = True
    doc.add_paragraph(
        "Partner zobowiazuje sie do zakonczenia wszystkich prac i przekazania wynikow prac "
        "Zamawiajacemu do dnia rozwiazania Umowy."
    )
    
    doc.add_paragraph()
    
    # §3
    p = doc.add_paragraph()
    p.add_run("Par. 3").bold = True
    doc.add_paragraph(
        "Zamawiajacy zobowiazuje sie do uregulowania wszystkich naleznosci wobec Partnera "
        "za uslugi wykonane do dnia rozwiazania Umowy, w terminie 14 dni od dnia otrzymania "
        "prawidlowo wystawionej faktury."
    )
    
    doc.add_paragraph()
    
    # §4
    p = doc.add_paragraph()
    p.add_run("Par. 4").bold = True
    doc.add_paragraph(
        "Strony oswiadczaja, iz z tytulu rozwiazania Umowy nie beda wzajemnie dochodzic "
        "zadnych roszczen, za wyjatkiem roszczen wynikajacych z par. 2 i 3 niniejszego Porozumienia."
    )
    
    if reason:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Par. 5 - Przyczyna rozwiazania").bold = True
        doc.add_paragraph(reason)
    
    if settlement_notes:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Uwagi dodatkowe:").bold = True
        doc.add_paragraph(settlement_notes)
    
    # §5/6
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Par. {'6' if reason else '5'}").bold = True
    doc.add_paragraph(
        "Porozumienie sporzadzono w dwoch jednobrzmiiacych egzemplarzach, po jednym dla kazdej ze Stron."
    )
    
    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "________________________"
    table.cell(0, 1).text = "________________________"
    table.cell(1, 0).text = "B2B.net S.A."
    table.cell(1, 1).text = contractor_name
    
    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    safe_number = contract_number.replace("/", "_")
    filename = f"Porozumienie_{safe_number}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    
    return output_path
