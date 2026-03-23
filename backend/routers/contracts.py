"""
Contracts router — all contract endpoints
"""
import os
from typing import List
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

import json as _json
from database import get_db, Contract, ContractAuditLog
from models.schemas import (
    ContractGenerateRequest,
    ContractModifyRequest,
    ContractRiskRequest,
    ContractResponse,
    ContractDetailResponse,
    GenerateResponse,
    ModifyResponse,
    ModifyDiff,
    RiskResponse,
    RiskItem,
)
from services.ai_service import standardize_it_area, modify_paragraph_text, assess_legal_risk
from services.docx_service import generate_contract_docx, get_paragraph_text, modify_paragraph_in_docx
from services.rag_service import search_legal_context
from services.contract_numbering import generate_contract_number

router = APIRouter()


@router.get("/stats/new")
def get_new_stats(db: Session = Depends(get_db)):
    """Stats for newly generated contracts only (not historical imports)."""
    from sqlalchemy import func
    query = db.query(Contract).filter(~Contract.number.startswith("H-"))
    total = query.count()
    avg_rate = query.with_entities(func.avg(Contract.rate)).scalar() or 0
    return {
        "total": total,
        "avg_rate": round(float(avg_rate), 2),
        "note": "Only PARAGRAF-generated contracts (excludes historical imports)",
    }


@router.get("/stats")
def get_stats(db: Session = Depends(get_db)):
    """Get contract statistics."""
    from sqlalchemy import func
    total = db.query(Contract).count()
    by_status = dict(db.query(Contract.status, func.count(Contract.id)).group_by(Contract.status).all())
    avg_rate = db.query(func.avg(Contract.rate)).scalar() or 0
    by_client = dict(db.query(Contract.client, func.count(Contract.id)).group_by(Contract.client).all())
    by_area = dict(db.query(Contract.it_area, func.count(Contract.id)).group_by(Contract.it_area).all())
    
    return {
        "total": total,
        "by_status": by_status,
        "avg_rate": round(float(avg_rate), 2),
        "by_client": by_client,
        "by_area": by_area,
    }


@router.post("/onboard")
def onboard_contractor(req: ContractGenerateRequest):
    """Full automated onboarding: generate → validate → do_podpisu → email preview."""
    from services.onboarding_service import run_onboarding
    return run_onboarding(req.model_dump())


@router.post("/generate", response_model=GenerateResponse)
def generate_contract(req: ContractGenerateRequest, db: Session = Depends(get_db)):
    """Generate a new B2B contract."""
    
    # 1. Standardize IT area
    it_area = standardize_it_area(req.obszar_it)
    
    # 2. Generate contract number
    contract_number = generate_contract_number(db)
    
    # 3. Generate DOCX
    file_path = generate_contract_docx(
        contract_number=contract_number,
        imie=req.imie,
        nazwisko=req.nazwisko,
        firma=req.firma,
        nip=req.nip,
        regon=req.regon or "",
        adres=req.adres,
        email=req.email,
        tel=req.tel,
        rola=req.rola,
        stawka=req.stawka,
        klient=req.klient,
        data_startu=req.data_startu,
        it_area=it_area,
        opis_projektu=req.opis_projektu or "",
        miasto_klienta=req.miasto_klienta,
    )
    
    # 4. Save to database
    contract = Contract(
        number=contract_number,
        contractor_name=f"{req.imie} {req.nazwisko}",
        contractor_firstname=req.imie,
        contractor_lastname=req.nazwisko,
        contractor_company=req.firma,
        contractor_nip=req.nip,
        contractor_regon=req.regon or "",
        contractor_address=req.adres,
        contractor_email=req.email,
        contractor_phone=req.tel,
        client=req.klient,
        role=req.rola,
        rate=req.stawka,
        it_area=it_area,
        it_area_raw=req.obszar_it,
        project_description=req.opis_projektu or "",
        client_city=req.miasto_klienta,
        start_date=req.data_startu,
        status="draft",
        file_path=file_path,
    )
    db.add(contract)
    db.commit()
    db.refresh(contract)
    
    # Audit log
    db.add(ContractAuditLog(
        contract_id=contract.id,
        action="created",
        details=_json.dumps({"number": contract_number, "client": req.klient, "rate": req.stawka}),
    ))
    db.commit()
    
    return GenerateResponse(
        success=True,
        contract=ContractDetailResponse(
            id=contract.id,
            number=contract.number,
            contractor_name=contract.contractor_name,
            contractor_firstname=contract.contractor_firstname,
            contractor_lastname=contract.contractor_lastname,
            contractor_company=contract.contractor_company,
            contractor_nip=contract.contractor_nip,
            contractor_regon=contract.contractor_regon,
            contractor_address=contract.contractor_address,
            contractor_email=contract.contractor_email,
            contractor_phone=contract.contractor_phone,
            client=contract.client,
            role=contract.role,
            rate=contract.rate,
            it_area=contract.it_area,
            it_area_raw=contract.it_area_raw,
            project_description=contract.project_description,
            client_city=contract.client_city,
            start_date=contract.start_date,
            status=contract.status,
            file_path=contract.file_path,
            created_at=contract.created_at,
        ),
        message=f"Umowa {contract_number} wygenerowana pomyślnie.",
        it_area_standardized=it_area,
    )


@router.post("/modify", response_model=ModifyResponse)
def modify_contract(req: ContractModifyRequest, db: Session = Depends(get_db)):
    """Modify contract paragraphs using AI."""
    
    contract = db.query(Contract).filter(Contract.id == req.contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    if not contract.file_path or not os.path.exists(contract.file_path):
        raise HTTPException(status_code=404, detail="Plik umowy nie znaleziony")
    
    diffs = []
    current_file = contract.file_path
    
    for change in req.zmiany:
        # Get original paragraph text
        original_text = get_paragraph_text(current_file, change.paragraf)
        if not original_text:
            original_text = f"[Paragraf {change.paragraf} — nie znaleziono automatycznie]"
        
        # AI modifies the text
        new_text = modify_paragraph_text(original_text, change.zmiana, change.paragraf)
        
        # Apply to DOCX
        new_file = modify_paragraph_in_docx(current_file, change.paragraf, new_text)
        current_file = new_file
        
        diffs.append(ModifyDiff(
            paragraf=change.paragraf,
            before=original_text,
            after=new_text,
        ))
    
    # Update DB
    contract.file_path = current_file
    contract.status = "modified"
    db.commit()
    
    return ModifyResponse(
        success=True,
        contract_id=req.contract_id,
        diffs=diffs,
        file_path=current_file,
        message=f"Wprowadzono {len(req.zmiany)} zmian(y) w umowie.",
    )


@router.post("/check-risks", response_model=RiskResponse)
def check_risks(req: ContractRiskRequest, db: Session = Depends(get_db)):
    """Check legal risks for proposed contract changes."""
    
    contract = db.query(Contract).filter(Contract.id == req.contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    risks = []
    risk_levels = {"green": 0, "yellow": 0, "red": 0}
    
    for change in req.zmiany:
        # Search RAG for legal context
        query = f"{change.paragraf}: {change.zmiana}"
        legal_context = search_legal_context(query)
        context_text = "\n".join(legal_context) if legal_context else ""
        
        # AI assesses risk
        assessment = assess_legal_risk(change.paragraf, change.zmiana, context_text)
        
        risk_level = assessment.get("ryzyko", "yellow")
        risk_levels[risk_level] = risk_levels.get(risk_level, 0) + 1
        
        risks.append(RiskItem(
            paragraf=change.paragraf,
            zmiana=change.zmiana,
            ryzyko=risk_level,
            uzasadnienie=assessment.get("uzasadnienie", ""),
            przepisy=assessment.get("przepisy", []),
        ))
    
    # Overall risk = worst
    if risk_levels["red"] > 0:
        overall = "red"
    elif risk_levels["yellow"] > 0:
        overall = "yellow"
    else:
        overall = "green"
    
    return RiskResponse(
        success=True,
        contract_id=req.contract_id,
        risks=risks,
        overall_risk=overall,
    )


@router.get("", response_model=List[ContractResponse])
def list_contracts(
    skip: int = 0,
    limit: int = 50,
    status: str = None,
    q: str = None,
    client: str = None,
    nip: str = None,
    db: Session = Depends(get_db),
):
    """List all contracts with optional filtering."""
    query = db.query(Contract)
    if status:
        query = query.filter(Contract.status == status)
    if q:
        from sqlalchemy import or_
        search = f"%{q}%"
        query = query.filter(or_(
            Contract.contractor_name.ilike(search),
            Contract.contractor_company.ilike(search),
            Contract.client.ilike(search),
            Contract.number.ilike(search),
            Contract.role.ilike(search),
        ))
    if client:
        query = query.filter(Contract.client.ilike(f"%{client}%"))
    if nip:
        query = query.filter(Contract.contractor_nip == nip.replace("-", "").replace(" ", ""))
    
    # Date range filter
    date_from: str = None
    date_to: str = None
    
    total_count = query.count()
    contracts = query.order_by(Contract.created_at.desc()).offset(skip).limit(limit).all()
    
    # Add total count to response header for pagination
    from fastapi import Response
    return contracts


@router.get("/{contract_id}", response_model=ContractDetailResponse)
def get_contract(contract_id: int, db: Session = Depends(get_db)):
    """Get contract details."""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    return contract


@router.get("/{contract_id}/download")
def download_contract(contract_id: int, format: str = "docx", db: Session = Depends(get_db)):
    """Download contract file. Format: docx (default) or pdf."""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    if not contract.file_path or not os.path.exists(contract.file_path):
        raise HTTPException(status_code=404, detail="Plik umowy nie znaleziony")
    
    if format == "pdf":
        # Try to convert to PDF using LibreOffice if available
        import subprocess
        pdf_path = contract.file_path.replace(".docx", ".pdf")
        if not os.path.exists(pdf_path):
            try:
                subprocess.run([
                    "soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", os.path.dirname(contract.file_path),
                    contract.file_path
                ], timeout=30, capture_output=True)
            except (FileNotFoundError, subprocess.TimeoutExpired):
                raise HTTPException(status_code=501, detail="Konwersja PDF niedostępna. Pobierz DOCX.")
        
        if os.path.exists(pdf_path):
            return FileResponse(
                path=pdf_path,
                filename=os.path.basename(pdf_path),
                media_type="application/pdf",
            )
        else:
            raise HTTPException(status_code=501, detail="Konwersja PDF nie powiodła się. Pobierz DOCX.")
    
    filename = os.path.basename(contract.file_path)
    return FileResponse(
        path=contract.file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.put("/{contract_id}")
def update_contract(contract_id: int, data: dict, db: Session = Depends(get_db)):
    """Update contract fields (rate, role, client, start_date, etc.)."""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    if contract.status not in ("draft", "modified"):
        raise HTTPException(status_code=400, detail="Można edytować tylko umowy robocze lub zmodyfikowane")
    
    editable_fields = ["role", "rate", "client", "start_date", "client_city",
                       "contractor_email", "contractor_phone", "project_description"]
    
    changed = {}
    for field in editable_fields:
        if field in data:
            old_val = getattr(contract, field, None)
            setattr(contract, field, data[field])
            changed[field] = {"from": old_val, "to": data[field]}
    
    if changed:
        db.add(ContractAuditLog(
            contract_id=contract_id,
            action="fields_updated",
            details=_json.dumps(changed),
        ))
        
        # Regenerate DOCX if key fields changed
        regen_fields = {"rate", "role", "client", "start_date", "client_city"}
        if regen_fields & set(changed.keys()) and contract.file_path:
            try:
                from services.docx_service import generate_contract_docx, _create_version
                _create_version(contract.file_path)
                new_path = generate_contract_docx(
                    contract_number=contract.number,
                    imie=contract.contractor_firstname,
                    nazwisko=contract.contractor_lastname,
                    firma=contract.contractor_company,
                    nip=contract.contractor_nip,
                    regon=contract.contractor_regon or "",
                    adres=contract.contractor_address,
                    email=contract.contractor_email,
                    tel=contract.contractor_phone,
                    rola=contract.role,
                    stawka=contract.rate,
                    klient=contract.client,
                    data_startu=contract.start_date,
                    it_area=contract.it_area,
                    opis_projektu=contract.project_description or "",
                    miasto_klienta=contract.client_city,
                )
                contract.file_path = new_path
            except Exception as e:
                pass  # Non-critical — keep old file
    
    db.commit()
    db.refresh(contract)
    return {"success": True, "updated_fields": list(changed.keys())}


@router.patch("/{contract_id}/status")
def update_status(contract_id: int, status: str, db: Session = Depends(get_db)):
    """Update contract status: draft → do_podpisu → aktywna → zakonczona."""
    valid_statuses = ["draft", "modified", "do_podpisu", "aktywna", "zakonczona", "anulowana"]
    if status not in valid_statuses:
        raise HTTPException(status_code=400, detail=f"Nieprawidłowy status. Dozwolone: {valid_statuses}")
    
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    # Pre-sign validation: check data completeness before activation
    if status == "aktywna" and contract.status == "do_podpisu":
        issues = []
        if not contract.contractor_nip or len(contract.contractor_nip.strip()) < 10:
            issues.append("Brak NIP kontraktora")
        if not contract.contractor_email:
            issues.append("Brak email kontraktora")
        if not contract.rate or contract.rate <= 0:
            issues.append("Brak stawki")
        if not contract.file_path or not os.path.exists(contract.file_path):
            issues.append("Brak wygenerowanego dokumentu DOCX")
        if issues:
            raise HTTPException(
                status_code=400,
                detail=f"Nie można aktywować umowy — brakujące dane: {', '.join(issues)}"
            )
    
    # Validate transition
    allowed_transitions = {
        "draft": ["do_podpisu", "modified", "anulowana"],
        "modified": ["do_podpisu", "draft", "anulowana"],
        "do_podpisu": ["aktywna", "draft", "anulowana"],
        "aktywna": ["zakonczona", "anulowana"],
        "zakonczona": [],
        "anulowana": ["draft"],
    }
    current = contract.status
    allowed = allowed_transitions.get(current, [])
    if status not in allowed and allowed:
        raise HTTPException(
            status_code=400,
            detail=f"Nie można zmienić statusu z '{current}' na '{status}'. Dozwolone: {allowed}"
        )
    
    old_status = contract.status
    contract.status = status
    
    # Log status notification for monitoring
    import logging
    logging.getLogger("paragraf").info(
        f"STATUS CHANGE: {contract.number} ({contract.contractor_name}) "
        f"{old_status} → {status} @ {contract.client}"
    )
    db.add(ContractAuditLog(
        contract_id=contract_id,
        action="status_changed",
        details=_json.dumps({"from": old_status, "to": status}),
    ))
    db.commit()
    return {"success": True, "message": f"Status zmieniony na: {status}"}


@router.post("/{contract_id}/tag")
def add_tag(contract_id: int, data: dict, db: Session = Depends(get_db)):
    """Add/remove tags on a contract."""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Not found")
    
    tag = data.get("tag", "")
    action = data.get("action", "add")  # "add" or "remove"
    
    existing = (getattr(contract, "tags", "") or "").split(",")
    existing = [t.strip() for t in existing if t.strip()]
    
    if action == "add" and tag not in existing:
        existing.append(tag)
    elif action == "remove" and tag in existing:
        existing.remove(tag)
    
    new_tags = ",".join(existing)
    try:
        contract.tags = new_tags
    except:
        pass
    
    db.commit()
    return {"success": True, "tags": existing}


@router.post("/{contract_id}/note")
def add_note(contract_id: int, data: dict, db: Session = Depends(get_db)):
    """Add a note to a contract."""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Not found")
    
    note = data.get("note", "")
    if not note:
        raise HTTPException(status_code=400, detail="Note text required")
    
    # Append to existing notes
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    existing = getattr(contract, "notes", "") or ""
    new_notes = f"{existing}\n[{timestamp}] {note}".strip()
    
    # Use raw SQL since column might not be in model yet
    try:
        contract.notes = new_notes
    except:
        db.execute(
            f"UPDATE contracts SET notes = ? WHERE id = ?",
            (new_notes, contract_id)
        )
    
    db.add(ContractAuditLog(
        contract_id=contract_id,
        action="note_added",
        details=_json.dumps({"note": note[:200]}),
    ))
    db.commit()
    
    return {"success": True, "notes": new_notes}


@router.post("/{contract_id}/regenerate")
def regenerate_docx(contract_id: int, db: Session = Depends(get_db)):
    """Regenerate DOCX file from current contract data (after field edits)."""
    from services.docx_service import generate_contract_docx, _create_version
    
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    # Backup old file
    if contract.file_path and os.path.exists(contract.file_path):
        _create_version(contract.file_path)
    
    # Regenerate
    new_path = generate_contract_docx(
        contract_number=contract.number,
        imie=contract.contractor_firstname,
        nazwisko=contract.contractor_lastname,
        firma=contract.contractor_company,
        nip=contract.contractor_nip,
        regon=contract.contractor_regon or "",
        adres=contract.contractor_address,
        email=contract.contractor_email,
        tel=contract.contractor_phone,
        rola=contract.role,
        stawka=contract.rate,
        klient=contract.client,
        data_startu=contract.start_date,
        it_area=contract.it_area,
        opis_projektu=contract.project_description or "",
        miasto_klienta=contract.client_city,
    )
    
    contract.file_path = new_path
    db.add(ContractAuditLog(
        contract_id=contract_id,
        action="regenerated",
        details=_json.dumps({"file": new_path}),
    ))
    db.commit()
    
    return {"success": True, "file_path": new_path, "message": f"DOCX regenerated for {contract.number}"}


@router.post("/{contract_id}/clone")
def clone_contract(contract_id: int, overrides: dict = {}, db: Session = Depends(get_db)):
    """Clone a contract with optional field overrides (new client, rate, role, etc.)."""
    from services.contract_numbering import generate_contract_number
    from services.docx_service import generate_contract_docx
    from services.ai_service import standardize_it_area
    
    original = db.query(Contract).filter(Contract.id == contract_id).first()
    if not original:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    new_number = generate_contract_number(db)
    
    # Apply overrides
    klient = overrides.get("klient", original.client)
    rola = overrides.get("rola", original.role)
    stawka = float(overrides.get("stawka", original.rate))
    data_startu = overrides.get("data_startu", original.start_date)
    miasto = overrides.get("miasto_klienta", original.client_city)
    obszar = overrides.get("obszar_it", original.it_area)
    
    if "obszar_it" in overrides:
        obszar = standardize_it_area(overrides["obszar_it"])
    
    # Generate new DOCX
    file_path = generate_contract_docx(
        contract_number=new_number,
        imie=original.contractor_firstname,
        nazwisko=original.contractor_lastname,
        firma=original.contractor_company,
        nip=original.contractor_nip,
        regon=original.contractor_regon or "",
        adres=original.contractor_address,
        email=original.contractor_email,
        tel=original.contractor_phone,
        rola=rola,
        stawka=stawka,
        klient=klient,
        data_startu=data_startu,
        it_area=obszar,
        opis_projektu=overrides.get("opis_projektu", original.project_description or ""),
        miasto_klienta=miasto,
    )
    
    new_contract = Contract(
        number=new_number,
        contractor_name=original.contractor_name,
        contractor_firstname=original.contractor_firstname,
        contractor_lastname=original.contractor_lastname,
        contractor_company=original.contractor_company,
        contractor_nip=original.contractor_nip,
        contractor_regon=original.contractor_regon,
        contractor_address=original.contractor_address,
        contractor_email=original.contractor_email,
        contractor_phone=original.contractor_phone,
        client=klient,
        role=rola,
        rate=stawka,
        it_area=obszar,
        it_area_raw=overrides.get("obszar_it", original.it_area_raw),
        project_description=overrides.get("opis_projektu", original.project_description),
        client_city=miasto,
        start_date=data_startu,
        status="draft",
        file_path=file_path,
    )
    db.add(new_contract)
    db.commit()
    db.refresh(new_contract)
    
    db.add(ContractAuditLog(
        contract_id=new_contract.id,
        action="cloned",
        details=_json.dumps({"from": contract_id, "overrides": overrides}),
    ))
    db.commit()
    
    return {
        "success": True,
        "contract_id": new_contract.id,
        "number": new_number,
        "message": f"Umowa sklonowana z {original.number} jako {new_number}",
    }


@router.post("/{contract_id}/duplicate")
def duplicate_contract(contract_id: int, db: Session = Depends(get_db)):
    """Duplicate a contract as new draft."""
    original = db.query(Contract).filter(Contract.id == contract_id).first()
    if not original:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    from services.contract_numbering import generate_contract_number
    new_number = generate_contract_number(db)
    
    new_contract = Contract(
        number=new_number,
        contractor_name=original.contractor_name,
        contractor_firstname=original.contractor_firstname,
        contractor_lastname=original.contractor_lastname,
        contractor_company=original.contractor_company,
        contractor_nip=original.contractor_nip,
        contractor_regon=original.contractor_regon,
        contractor_address=original.contractor_address,
        contractor_email=original.contractor_email,
        contractor_phone=original.contractor_phone,
        client=original.client,
        role=original.role,
        rate=original.rate,
        it_area=original.it_area,
        it_area_raw=original.it_area_raw,
        project_description=original.project_description,
        client_city=original.client_city,
        start_date=original.start_date,
        status="draft",
        file_path=original.file_path,
    )
    db.add(new_contract)
    db.commit()
    db.refresh(new_contract)
    return {"success": True, "contract_id": new_contract.id, "number": new_number}


@router.post("/{contract_id}/terminate")
def terminate_contract(contract_id: int, data: dict, db: Session = Depends(get_db)):
    """Generate mutual termination agreement and close contract."""
    from services.termination_service import generate_termination
    
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    termination_date = data.get("termination_date", "")
    reason = data.get("reason", "")
    notes = data.get("notes", "")
    
    if not termination_date:
        raise HTTPException(status_code=400, detail="Podaj datę rozwiązania")
    
    file_path = generate_termination(
        contract_number=contract.number,
        contractor_name=contract.contractor_name,
        contractor_company=contract.contractor_company,
        contractor_nip=contract.contractor_nip,
        termination_date=termination_date,
        reason=reason,
        settlement_notes=notes,
    )
    
    # Update contract status
    contract.status = "zakonczona"
    contract.end_date = termination_date.replace(".", "-") if "." in termination_date else termination_date
    
    db.add(ContractAuditLog(
        contract_id=contract_id,
        action="terminated",
        details=_json.dumps({"date": termination_date, "reason": reason, "file": file_path}),
    ))
    db.commit()
    
    return {
        "success": True,
        "file_path": file_path,
        "message": f"Porozumienie o rozwiązaniu umowy {contract.number} z dniem {termination_date}",
    }


@router.post("/{contract_id}/annex")
def create_annex(contract_id: int, data: dict, db: Session = Depends(get_db)):
    """Generate an annex (aneks) for contract modifications."""
    from services.annex_service import generate_annex
    
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    annex_type = data.get("type", "other")
    changes = data.get("changes", [])
    effective_date = data.get("effective_date", "")
    notes = data.get("notes", "")
    
    if not changes:
        raise HTTPException(status_code=400, detail="Brak zmian do aneksu")
    
    file_path = generate_annex(
        contract_number=contract.number,
        contractor_name=contract.contractor_name,
        contractor_company=contract.contractor_company,
        contractor_nip=contract.contractor_nip,
        contractor_address=contract.contractor_address,
        annex_type=annex_type,
        changes=changes,
        effective_date=effective_date,
        notes=notes,
    )
    
    # Update contract fields if rate/role changed
    for change in changes:
        if change.get("field") == "stawka" and change.get("new"):
            try:
                contract.rate = float(change["new"])
            except ValueError:
                pass
        if change.get("field") == "rola" and change.get("new"):
            contract.role = change["new"]
        if change.get("field") == "klient" and change.get("new"):
            contract.client = change["new"]
    
    # Audit
    db.add(ContractAuditLog(
        contract_id=contract_id,
        action="annex_created",
        details=_json.dumps({"type": annex_type, "changes": changes, "file": file_path}),
    ))
    db.commit()
    
    return {
        "success": True,
        "file_path": file_path,
        "message": f"Aneks wygenerowany dla umowy {contract.number}",
    }


@router.post("/{contract_id}/send-welcome-email")
def send_welcome_email(contract_id: int, db: Session = Depends(get_db)):
    """Send welcome email to contractor via Microsoft Graph."""
    from services.email_service import generate_welcome_email
    from services.msgraph_service import send_email, is_configured
    
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    if not contract.contractor_email:
        raise HTTPException(status_code=400, detail="Brak adresu email kontraktora")
    
    email = generate_welcome_email(contract)
    
    if not is_configured():
        return {
            "success": False,
            "preview": email,
            "error": "MS365 token not configured. Email preview shown.",
        }
    
    result = send_email(
        to=email["to"],
        subject=email["subject"],
        body_html=email["body_html"],
    )
    
    if result["success"]:
        db.add(ContractAuditLog(
            contract_id=contract_id,
            action="email_sent",
            details=_json.dumps({"to": email["to"], "subject": email["subject"][:50]}),
        ))
        db.commit()
    
    return result


@router.post("/{contract_id}/full-review")
def full_review(contract_id: int, db: Session = Depends(get_db)):
    """Run full AI review of a contract (uses Sonnet — costs more)."""
    from services.ai_service import full_contract_review
    
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    if not contract.file_path or not os.path.exists(contract.file_path):
        raise HTTPException(status_code=404, detail="Plik umowy nie znaleziony")
    
    # Extract text from DOCX
    from docx import Document as DocxDocument
    doc = DocxDocument(contract.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    
    review = full_contract_review(full_text, contract.contractor_name, contract.client)
    
    # Audit
    db.add(ContractAuditLog(
        contract_id=contract_id,
        action="full_review",
        details=_json.dumps({"score": review.get("overall_score", 0)}),
    ))
    db.commit()
    
    return review


@router.get("/{contract_id}/email-preview")
def email_preview(contract_id: int, db: Session = Depends(get_db)):
    """Generate welcome email preview for contractor."""
    from services.email_service import generate_welcome_email
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    return generate_welcome_email(contract)


@router.get("/{contract_id}/markdown")
def contract_as_markdown(contract_id: int, db: Session = Depends(get_db)):
    """Export contract data as markdown."""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Not found")
    
    md = f"""# Umowa {contract.number}

## Kontrahent
- **Imię i nazwisko:** {contract.contractor_name}
- **Firma:** {contract.contractor_company}
- **NIP:** {contract.contractor_nip}
- **Adres:** {contract.contractor_address}
- **Email:** {contract.contractor_email}
- **Telefon:** {contract.contractor_phone}

## Projekt
- **Klient:** {contract.client}
- **Rola:** {contract.role}
- **Stawka:** {contract.rate} PLN/h netto
- **Obszar IT:** {contract.it_area}
- **Miasto:** {contract.client_city}
- **Data startu:** {contract.start_date}

## Status
- **Status:** {contract.status}
- **Data utworzenia:** {contract.created_at}

---
*PARAGRAF · B2B.net S.A.*
"""
    return {"markdown": md, "number": contract.number}


@router.get("/{contract_id}/diff-with-template")
def diff_with_template(contract_id: int, db: Session = Depends(get_db)):
    """Compare contract with original template — show what changed."""
    from docx import Document as DocxDocument
    
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract or not contract.file_path or not os.path.exists(contract.file_path):
        raise HTTPException(status_code=404, detail="Contract or file not found")
    
    template_path = os.path.expanduser("~/clawd/data/legal/marta_templates/Umowa_B2B_(draft_2026)od_02.03.2026.docx")
    
    contract_doc = DocxDocument(contract.file_path)
    template_doc = DocxDocument(template_path)
    
    contract_paras = [p.text.strip() for p in contract_doc.paragraphs if p.text.strip()]
    template_paras = [p.text.strip() for p in template_doc.paragraphs if p.text.strip()]
    
    # Find differences
    changed = []
    for i, (cp, tp) in enumerate(zip(contract_paras[:50], template_paras[:50])):
        if cp != tp and len(cp) > 20 and len(tp) > 20:
            # Check similarity
            common = sum(1 for a, b in zip(cp, tp) if a == b)
            similarity = common / max(len(cp), len(tp), 1) * 100
            if similarity < 95:  # Only show real differences
                changed.append({
                    "index": i,
                    "template": tp[:200],
                    "contract": cp[:200],
                    "similarity_pct": round(similarity, 1),
                })
    
    return {
        "contract_id": contract_id,
        "number": contract.number,
        "template_paragraphs": len(template_paras),
        "contract_paragraphs": len(contract_paras),
        "differences": changed[:15],
        "total_differences": len(changed),
    }


@router.get("/{contract_id}/similar")
def find_similar(contract_id: int, db: Session = Depends(get_db)):
    """Find contracts similar to this one (same area + client)."""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    similar = db.query(Contract).filter(
        Contract.id != contract_id,
        Contract.client == contract.client,
        Contract.it_area == contract.it_area,
    ).order_by(Contract.created_at.desc()).limit(10).all()
    
    if not similar:
        # Relax: just same client
        similar = db.query(Contract).filter(
            Contract.id != contract_id,
            Contract.client == contract.client,
        ).order_by(Contract.created_at.desc()).limit(5).all()
    
    return {
        "contract_id": contract_id,
        "reference": {"client": contract.client, "area": contract.it_area, "rate": contract.rate},
        "similar": [
            {
                "id": c.id, "number": c.number, "contractor": c.contractor_name,
                "role": c.role, "rate": c.rate, "status": c.status,
                "rate_diff": round(c.rate - contract.rate, 2),
            }
            for c in similar
        ],
        "avg_rate_similar": round(
            sum(c.rate for c in similar if c.rate > 0) / max(1, sum(1 for c in similar if c.rate > 0)), 2
        ),
    }


@router.get("/{contract_id}/summary")
def get_summary(contract_id: int, db: Session = Depends(get_db)):
    """Get AI-generated summary of the contract."""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    import anthropic
    client = anthropic.Anthropic()
    
    msg = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=300,
        messages=[{"role": "user", "content": f"""Wygeneruj krótkie podsumowanie umowy B2B w 3-4 zdaniach po polsku:

Numer: {contract.number}
Kontrahent: {contract.contractor_name} ({contract.contractor_company})
NIP: {contract.contractor_nip}
Klient: {contract.client}
Rola: {contract.role}
Stawka: {contract.rate} PLN/h
Obszar IT: {contract.it_area}
Data startu: {contract.start_date}
Status: {contract.status}
Miasto: {contract.client_city}

Podsumuj kto, co, dla kogo, za ile. Bądź zwięzły."""}],
    )
    
    return {
        "contract_id": contract_id,
        "summary": msg.content[0].text.strip(),
    }


@router.get("/{contract_id}/versions")
def get_versions(contract_id: int, db: Session = Depends(get_db)):
    """Get list of DOCX versions for a contract."""
    import glob
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract or not contract.file_path:
        return []
    
    versions_dir = os.path.join(os.path.dirname(contract.file_path), "versions")
    base = os.path.basename(contract.file_path).replace(".docx", "")
    
    versions = []
    if os.path.exists(versions_dir):
        files = sorted(glob.glob(os.path.join(versions_dir, f"{base}_v*.docx")))
        for f in files:
            stat = os.stat(f)
            versions.append({
                "filename": os.path.basename(f),
                "path": f,
                "size_kb": round(stat.st_size / 1024, 1),
                "created": stat.st_mtime,
            })
    
    # Add current version
    if os.path.exists(contract.file_path):
        stat = os.stat(contract.file_path)
        versions.append({
            "filename": os.path.basename(contract.file_path) + " (aktualna)",
            "path": contract.file_path,
            "size_kb": round(stat.st_size / 1024, 1),
            "created": stat.st_mtime,
        })
    
    return versions


@router.get("/{contract_id}/history")
def get_history(contract_id: int, db: Session = Depends(get_db)):
    """Get audit log history for a contract."""
    logs = db.query(ContractAuditLog).filter(
        ContractAuditLog.contract_id == contract_id
    ).order_by(ContractAuditLog.created_at.desc()).all()
    
    return [
        {
            "id": log.id,
            "action": log.action,
            "details": _json.loads(log.details) if log.details else {},
            "user": log.user,
            "created_at": log.created_at.isoformat(),
        }
        for log in logs
    ]


@router.get("/{contract_id}/html")
def get_contract_html(contract_id: int, db: Session = Depends(get_db)):
    """Get contract as HTML for browser preview/print."""
    from services.html_service import docx_to_html
    from fastapi.responses import HTMLResponse
    
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    if not contract.file_path or not os.path.exists(contract.file_path):
        raise HTTPException(status_code=404, detail="Plik nie znaleziony")
    
    html = docx_to_html(contract.file_path)
    return HTMLResponse(content=html)


@router.get("/{contract_id}/check-quality")
def check_quality(contract_id: int, db: Session = Depends(get_db)):
    """Check if contract DOCX has unfilled placeholders."""
    import re
    from docx import Document as DocxDocument
    
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    if not contract.file_path or not os.path.exists(contract.file_path):
        raise HTTPException(status_code=404, detail="Plik umowy nie znaleziony")
    
    doc = DocxDocument(contract.file_path)
    issues = []
    
    contractor_name = contract.contractor_name
    for para in doc.paragraphs:
        text = para.text
        # Skip signature lines (they should have blanks)
        is_signature_line = (
            "Podpis" in text or 
            "B2B.NET S.A." in text or 
            contractor_name in text or
            text.count("_") > 10  # pure separator line
        )
        # Check for unfilled placeholders
        if re.search(r'_{5,}', text) and not is_signature_line:
            issues.append({"type": "blank_field", "text": text[:100]})
        if re.search(r'…{5,}', text) and "B2B.net" not in text and not is_signature_line:
            issues.append({"type": "unfilled_dots", "text": text[:100]})
        if "[PEŁNA NAZWA" in text or "[NUMER]" in text:
            issues.append({"type": "unfilled_bracket", "text": text[:100]})
    
    # Check tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if "Wpisujemy" in text:
                    issues.append({"type": "instruction_text", "text": text[:100]})
    
    return {
        "contract_id": contract_id,
        "number": contract.number,
        "issues_count": len(issues),
        "issues": issues,
        "status": "ok" if not issues else "has_issues",
    }


@router.post("/batch-onboard")
def batch_onboard(contractors: List[dict], db: Session = Depends(get_db)):
    """Onboard multiple contractors at once."""
    from services.onboarding_service import run_onboarding
    results = []
    for i, data in enumerate(contractors):
        try:
            result = run_onboarding(data)
            results.append({"index": i, "success": True, "number": result.get("number"), "id": result.get("contract_id")})
        except Exception as e:
            results.append({"index": i, "success": False, "error": str(e)})
    
    success_count = sum(1 for r in results if r["success"])
    return {"total": len(contractors), "success": success_count, "failed": len(contractors) - success_count, "results": results}


@router.post("/batch-generate")
def batch_generate(contracts: List[dict], db: Session = Depends(get_db)):
    """Generate multiple contracts at once. Each item is a ContractGenerateRequest."""
    results = []
    errors = []
    
    for i, data in enumerate(contracts):
        try:
            # Reuse the generate logic
            from models.schemas import ContractGenerateRequest
            req = ContractGenerateRequest(**data)
            
            it_area = standardize_it_area(req.obszar_it)
            contract_number = generate_contract_number(db)
            
            from services.docx_service import generate_contract_docx
            file_path = generate_contract_docx(
                contract_number=contract_number,
                imie=req.imie, nazwisko=req.nazwisko,
                firma=req.firma, nip=req.nip,
                regon=req.regon or "", adres=req.adres,
                email=req.email, tel=req.tel,
                rola=req.rola, stawka=req.stawka,
                klient=req.klient, data_startu=req.data_startu,
                it_area=it_area,
                opis_projektu=req.opis_projektu or "",
                miasto_klienta=req.miasto_klienta,
            )
            
            contract = Contract(
                number=contract_number,
                contractor_name=f"{req.imie} {req.nazwisko}",
                contractor_firstname=req.imie,
                contractor_lastname=req.nazwisko,
                contractor_company=req.firma,
                contractor_nip=req.nip,
                contractor_regon=req.regon or "",
                contractor_address=req.adres,
                contractor_email=req.email,
                contractor_phone=req.tel,
                client=req.klient,
                role=req.rola,
                rate=req.stawka,
                it_area=it_area,
                it_area_raw=req.obszar_it,
                project_description=req.opis_projektu or "",
                client_city=req.miasto_klienta,
                start_date=req.data_startu,
                status="draft",
                file_path=file_path,
            )
            db.add(contract)
            db.flush()
            results.append({"index": i, "number": contract_number, "id": contract.id})
        except Exception as e:
            errors.append({"index": i, "error": str(e)})
    
    db.commit()
    return {"generated": len(results), "errors": len(errors), "results": results, "error_details": errors}


@router.post("/bulk/terminate-client")
def bulk_terminate_client(client: str, termination_date: str, db: Session = Depends(get_db)):
    """Terminate all active contracts for a specific client (client leaves)."""
    contracts = db.query(Contract).filter(
        Contract.client.ilike(f"%{client}%"),
        Contract.status == "aktywna",
    ).all()
    
    if not contracts:
        raise HTTPException(status_code=404, detail=f"Brak aktywnych umów dla klienta: {client}")
    
    count = 0
    for c in contracts:
        c.status = "zakonczona"
        c.end_date = termination_date
        db.add(ContractAuditLog(
            contract_id=c.id,
            action="bulk_terminated",
            details=_json.dumps({"client": client, "date": termination_date}),
        ))
        count += 1
    
    db.commit()
    return {"success": True, "terminated": count, "client": client, "date": termination_date}


@router.post("/bulk/update")
def bulk_update(updates: List[dict], db: Session = Depends(get_db)):
    """Bulk update contract fields. Each item: {id, field, value}."""
    allowed_fields = ["contractor_nip", "rate", "contractor_email", "contractor_phone", 
                       "contractor_address", "contractor_company", "status"]
    
    updated = 0
    errors = []
    for item in updates:
        cid = item.get("id")
        field = item.get("field")
        value = item.get("value")
        
        if field not in allowed_fields:
            errors.append(f"Field '{field}' not allowed")
            continue
        
        contract = db.query(Contract).filter(Contract.id == cid).first()
        if not contract:
            errors.append(f"Contract {cid} not found")
            continue
        
        setattr(contract, field, value)
        updated += 1
    
    db.commit()
    return {"success": True, "updated": updated, "errors": errors}


@router.post("/bulk/status")
def bulk_status_update(contract_ids: List[int], status: str, db: Session = Depends(get_db)):
    """Update status for multiple contracts at once."""
    valid_statuses = ["draft", "modified", "do_podpisu", "aktywna", "zakonczona", "anulowana"]
    if status not in valid_statuses:
        raise HTTPException(status_code=400, detail=f"Nieprawidłowy status: {status}")
    
    updated = 0
    errors = []
    for cid in contract_ids:
        contract = db.query(Contract).filter(Contract.id == cid).first()
        if not contract:
            errors.append(f"Umowa {cid} nie znaleziona")
            continue
        contract.status = status
        updated += 1
    
    db.commit()
    return {"success": True, "updated": updated, "errors": errors}


@router.post("/merge")
def merge_contracts(source_id: int, target_id: int, db: Session = Depends(get_db)):
    """Merge source contract data INTO target (updates target, deletes source)."""
    source = db.query(Contract).filter(Contract.id == source_id).first()
    target = db.query(Contract).filter(Contract.id == target_id).first()
    
    if not source or not target:
        raise HTTPException(status_code=404, detail="Source or target not found")
    
    # Copy missing fields from source to target
    fields_to_merge = [
        "contractor_nip", "contractor_regon", "contractor_email", 
        "contractor_phone", "contractor_address", "contractor_company",
    ]
    
    merged_fields = []
    for field in fields_to_merge:
        src_val = getattr(source, field, "")
        tgt_val = getattr(target, field, "")
        if src_val and not tgt_val:
            setattr(target, field, src_val)
            merged_fields.append(field)
    
    # Audit
    db.add(ContractAuditLog(
        contract_id=target_id,
        action="merged",
        details=_json.dumps({"from": source_id, "from_number": source.number, "merged_fields": merged_fields}),
    ))
    
    # Mark source as merged/anulowana
    source.status = "anulowana"
    source.project_description = f"MERGED INTO {target.number} (ID:{target_id})"
    
    db.commit()
    
    return {
        "success": True,
        "target_id": target_id,
        "merged_fields": merged_fields,
        "message": f"Umowa {source.number} scalona z {target.number}",
    }


@router.delete("/{contract_id}")
def delete_contract(contract_id: int, db: Session = Depends(get_db)):
    """Delete a contract. Only draft or cancelled contracts can be deleted."""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Umowa nie znaleziona")
    
    if contract.status not in ("draft", "anulowana"):
        raise HTTPException(
            status_code=400,
            detail=f"Nie można usunąć umowy ze statusem '{contract.status}'. Tylko robocze i anulowane."
        )
    
    # Also delete audit log entries
    db.query(ContractAuditLog).filter(ContractAuditLog.contract_id == contract_id).delete()
    db.delete(contract)
    db.commit()
    return {"success": True, "message": "Umowa usunięta"}
